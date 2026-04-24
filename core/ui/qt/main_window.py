from __future__ import annotations

import logging
import re
import sqlite3
import uuid
from datetime import date, datetime
from decimal import ROUND_HALF_UP, Decimal, InvalidOperation
from pathlib import Path

from PySide6.QtCore import QMarginsF, QRect, Qt, QUrl
from PySide6.QtGui import (
    QColor,
    QDesktopServices,
    QFont,
    QPageLayout,
    QPageSize,
    QPainter,
    QPdfWriter,
    QPixmap,
)
from PySide6.QtWidgets import (
    QDockWidget,
    QFileDialog,
    QFrame,
    QHBoxLayout,
    QLabel,
    QMainWindow,
    QMessageBox,
    QPushButton,
    QStatusBar,
    QTabWidget,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from core.app_context import BudgetPalContext
from core.build_info import load_build_info
from core.domain import TransactionInput, TransferInput
from core.importers.subtracker_view import SubTrackerIntegrationError
from core.importers.xlsx_transactions import XLSXTransactionImporter
from core.path_registry import BudgetPalPathRegistry
from core.report_catalog import report_rows, report_type_lookup
from core.services.help_service import HelpService
from core.services.reporting import ReportingService
from core.settings import get_settings_manager
from core.ui.qt.budget_category_definitions_dialog import BudgetCategoryDefinitionsDialog
from core.ui.qt.income_definitions_dialog import IncomeDefinitionsDialog
from core.ui.qt.recurring_definitions_dialog import RecurringDefinitionsDialog
from core.ui.qt.settings_dialog import SettingsDialog
from core.ui.qt.sub_payments_dialog import SubPaymentsDialog
from core.ui.qt.tabs.accounts import AccountsTab
from core.ui.qt.tabs.bills import BillsTab
from core.ui.qt.tabs.budget_month import BudgetMonthTab
from core.ui.qt.tabs.dashboard import DashboardTab
from core.ui.qt.tabs.income import IncomeTab
from core.ui.qt.tabs.reports import ReportsTab
from core.ui.qt.tabs.transactions import TransactionsTab
from core.ui.qt.tabs.transfers import TransfersTab


class BudgetPalWindow(QMainWindow):
    BACKUP_KEEP_COUNT = 5
    LOG_LEVEL_COLORS = {
        "DEBUG": "#6B7280",
        "INFO": "#0F766E",
        "WARNING": "#B45309",
        "ERROR": "#B91C1C",
        "CRITICAL": "#7F1D1D",
        "DEFAULT": "#111827",
    }
    BUTTON_STYLESHEET = """
    QPushButton {
        background-color: #2563EB;
        color: #FFFFFF;
        border: 1px solid #1D4ED8;
        border-radius: 4px;
        padding: 4px 10px;
        font-weight: 600;
    }
    QPushButton:hover {
        background-color: #1D4ED8;
    }
    QPushButton:pressed {
        background-color: #1E40AF;
    }
    QPushButton:disabled {
        background-color: #93C5FD;
        color: #E5E7EB;
        border: 1px solid #93C5FD;
    }
    """

    def __init__(self, context: BudgetPalContext, logger: logging.Logger, log_emitter) -> None:
        super().__init__()
        self.context = context
        self.logger = logger
        self.log_emitter = log_emitter
        self.reporting_service = ReportingService(context.db)
        self.help_service = HelpService()
        self.setStyleSheet(self.BUTTON_STYLESHEET)
        self.account_alias_to_id: dict[str, int] = {}
        self.account_id_to_alias: dict[int, str] = {}
        self.account_id_to_type: dict[int, str] = {}
        self.account_ids_by_type: dict[str, list[int]] = {}
        self._suppress_selection_autoload = False
        self._suppress_type_defaults = False
        self.sub_payments_dialog: SubPaymentsDialog | None = None
        self.bill_definitions_dialog: RecurringDefinitionsDialog | None = None
        self.income_definitions_dialog: IncomeDefinitionsDialog | None = None
        self.budget_definitions_dialog: BudgetCategoryDefinitionsDialog | None = None
        self.bills_sort_key = "payment_due"
        self.income_sort_key = "description"
        self.budget_sort_key = "category"
        self._editing_bill_source_system: str | None = None
        self._bills_dirty_by_month: dict[str, bool] = {}
        self._subscriptions_dirty_by_month: dict[str, bool] = {}
        self._income_dirty_by_month: dict[str, bool] = {}
        self._budget_dirty_by_month: dict[str, bool] = {}
        self._loading_dashboard_starting_balance = False
        self.editing_transfer_group_id: str | None = None
        self.editing_transfer_source_system: str | None = None

        self.setWindowTitle("BudgetPal")
        configured_width = int(self.context.settings["ui"]["window"].get("width", 1240))
        configured_height = int(self.context.settings["ui"]["window"].get("height", 820))
        self.resize(
            max(800, int(round(configured_width * 1.25))),
            configured_height,
        )

        self.tabs = QTabWidget()
        self.dashboard_tab = DashboardTab()
        self.transactions_tab = TransactionsTab()
        self.transfers_tab = TransfersTab()
        self.budget_tab = BudgetMonthTab()
        self.bills_tab = BillsTab()
        self.income_tab = IncomeTab()
        self.accounts_tab = AccountsTab()
        self.reports_tab = ReportsTab()

        self.tabs.addTab(self.dashboard_tab, "Dashboard")
        self.tabs.addTab(self.transactions_tab, "Transactions")
        self.tabs.addTab(self.transfers_tab, "Transfers")
        self.tabs.addTab(self.income_tab, "Income")
        self.tabs.addTab(self.bills_tab, "Bills")
        self.tabs.addTab(self.budget_tab, "Budget Allocations")
        self.tabs.addTab(self.accounts_tab, "Accounts")
        self.tabs.addTab(self.reports_tab, "Reports")
        self._init_central_layout()

        self._init_log_dock()
        self.setStatusBar(QStatusBar())
        today = date.today()
        self.transactions_view_year = today.year
        self.transactions_view_month = today.month
        self.budget_view_year = today.year
        self.budget_view_month = today.month
        self.transfers_view_year = today.year
        self.transfers_view_month = today.month
        self.bills_view_year = today.year
        self.bills_view_month = today.month
        self.income_view_year = today.year
        self.income_view_month = today.month
        self.accounts_view_year = today.year
        self.accounts_view_month = today.month

        self._populate_month_selectors()
        self._wire_events()
        self._refresh_dashboard_month_filter(
            preferred_month=f"{today.year}-{today.month:02d}"
        )
        self._refresh_transactions_month_filter(
            preferred_month=f"{self.transactions_view_year}-{self.transactions_view_month:02d}"
        )
        self._refresh_transfers_month_filter(
            preferred_month=f"{self.transfers_view_year}-{self.transfers_view_month:02d}"
        )
        self._refresh_budget_month_filter(
            preferred_month=f"{self.budget_view_year}-{self.budget_view_month:02d}"
        )
        self._refresh_bills_month_filter(
            preferred_month=f"{self.bills_view_year}-{self.bills_view_month:02d}"
        )
        self._refresh_income_month_filter(
            preferred_month=f"{self.income_view_year}-{self.income_view_month:02d}"
        )
        self._refresh_accounts_month_filter(
            preferred_month=f"{self.accounts_view_year}-{self.accounts_view_month:02d}"
        )
        self._sync_reports_period_from_dashboard(
            preferred_month=f"{today.year}-{today.month:02d}"
        )
        self._restore_reports_table_column_widths()
        self._load_static_reports()
        self._refresh_transaction_form_choices()
        self._refresh_transfer_form_choices()
        self._clear_transaction_form()
        self.new_transfer_form()
        self._refresh_budget_form_choices()
        self._refresh_bill_form_choices()
        self._refresh_income_form_choices()
        self.new_budget_form()
        self.new_bill_form()
        self.new_income_form()
        self.refresh_transactions()
        self.refresh_transfers()
        self.refresh_budget_allocations()
        self.refresh_bills()
        self.refresh_income()
        self.refresh_accounts()
        self.refresh_dashboard()

        self.log_emitter.message.connect(self.append_log_message)
        self.logger.info("BudgetPal UI initialized")

    def _init_log_dock(self) -> None:
        dock = QDockWidget("Activity Log", self)
        dock.setAllowedAreas(Qt.BottomDockWidgetArea)

        holder = QWidget()
        layout = QVBoxLayout(holder)

        controls = QHBoxLayout()
        controls.addWidget(QLabel("Operations, errors, and status events"))
        controls.addStretch(1)
        clear_button = QPushButton("Clear")
        clear_button.clicked.connect(self._clear_log)
        controls.addWidget(clear_button)
        layout.addLayout(controls)

        self.log_area = QTextEdit()
        self.log_area.setReadOnly(True)
        self.log_area.document().setMaximumBlockCount(1000)
        layout.addWidget(self.log_area)

        dock.setWidget(holder)
        self.addDockWidget(Qt.BottomDockWidgetArea, dock)

    def _init_central_layout(self) -> None:
        root = QWidget()
        root_layout = QVBoxLayout(root)
        root_layout.setContentsMargins(2, 8, 8, 8)
        root_layout.setSpacing(8)

        header_frame = QFrame()
        header_layout = QHBoxLayout(header_frame)
        header_layout.setContentsMargins(24, 4, 10, 4)
        header_layout.setSpacing(24)

        self.logo_label = QLabel()
        self.logo_label.setAlignment(Qt.AlignCenter)
        self.logo_label.setFixedSize(110, 110)
        logo_path = BudgetPalPathRegistry.logo_image_file()
        if logo_path and logo_path.exists():
            pixmap = QPixmap(str(logo_path))
            if not pixmap.isNull():
                self.logo_label.setPixmap(
                    pixmap.scaled(
                        self.logo_label.size(),
                        Qt.KeepAspectRatio,
                        Qt.SmoothTransformation,
                    )
                )
            else:
                self.logo_label.setText("BudgetPal")
        else:
            self.logo_label.setText("BudgetPal")
        header_layout.addWidget(self.logo_label)

        title_block = QWidget()
        title_layout = QVBoxLayout(title_block)
        title_layout.setContentsMargins(0, 0, 0, 0)
        title_layout.setSpacing(0)

        self.app_title_label = QLabel("BudgetPal")
        # Let the system palette choose text color so contrast adapts to light/dark themes.
        self.app_title_label.setStyleSheet("font-size: 30px; font-weight: 700;")

        title_layout.addWidget(self.app_title_label)

        header_layout.addWidget(title_block)
        header_layout.addStretch(1)

        self.current_tab_header_label = QLabel(self.tabs.tabText(self.tabs.currentIndex()))
        self.current_tab_header_label.setAlignment(Qt.AlignCenter)
        self.current_tab_header_label.setStyleSheet("font-size: 20px; font-weight: 600;")
        header_layout.addWidget(self.current_tab_header_label, 0, Qt.AlignCenter)
        header_layout.addStretch(1)

        self.help_button = QPushButton("Help")
        self.settings_button = QPushButton("Settings")
        self.about_button = QPushButton("About")
        self.exit_button = QPushButton("Exit")
        header_layout.addWidget(self.help_button)
        header_layout.addWidget(self.settings_button)
        header_layout.addWidget(self.about_button)
        header_layout.addWidget(self.exit_button)

        root_layout.addWidget(header_frame)
        root_layout.addWidget(self.tabs, 1)
        self.setCentralWidget(root)

    def _populate_month_selectors(self) -> None:
        current = date.today()
        month_labels = []
        for offset in range(-3, 10):
            month = current.month + offset
            year = current.year + (month - 1) // 12
            month = ((month - 1) % 12) + 1
            month_labels.append(f"{year}-{month:02d}")

        selectors = [
            self.budget_tab.month_picker,
        ]
        for selector in selectors:
            selector.clear()
            selector.addItems(month_labels)
            selector.setCurrentText(f"{current.year}-{current.month:02d}")

    @staticmethod
    def _rolling_month_labels(months_back: int = 12, months_forward: int = 12) -> list[str]:
        current = date.today()
        labels: list[str] = []
        for offset in range(-months_back, months_forward + 1):
            month = current.month + offset
            year = current.year + (month - 1) // 12
            month = ((month - 1) % 12) + 1
            labels.append(f"{year}-{month:02d}")
        return labels

    def _wire_events(self) -> None:
        self.bills_tab.refresh_subtracker_button.clicked.connect(self.refresh_subscriptions)
        self.dashboard_tab.starting_balance_input.editingFinished.connect(self.save_dashboard_starting_balance)
        self.transactions_tab.import_button.clicked.connect(self.import_transactions_xlsx)
        self.transactions_tab.sub_payments_button.clicked.connect(self.show_sub_payments_dialog)
        self.transactions_tab.add_button.clicked.connect(self.add_transaction)
        self.transactions_tab.save_button.clicked.connect(self.save_transaction)
        self.transactions_tab.delete_button.clicked.connect(self.delete_selected_transaction)
        self.transactions_tab.expenses_table.clicked.connect(self._on_expense_row_clicked)
        self.transactions_tab.income_table.clicked.connect(self._on_income_row_clicked)
        self.transactions_tab.expenses_table.selectionModel().selectionChanged.connect(
            lambda *_: self.on_transaction_selection_changed()
        )
        self.transactions_tab.income_table.selectionModel().selectionChanged.connect(
            lambda *_: self.on_transaction_selection_changed()
        )
        self.transactions_tab.month_filter.currentTextChanged.connect(self.on_transactions_month_changed)
        self.transactions_tab.expense_radio.toggled.connect(lambda checked: self._on_type_changed("expense", checked))
        self.transactions_tab.income_radio.toggled.connect(lambda checked: self._on_type_changed("income", checked))
        self.transfers_tab.transfer_save_button.clicked.connect(self.save_transfer_from_form)
        self.transfers_tab.transfer_delete_button.clicked.connect(self.delete_selected_transfer)
        self.transfers_tab.transfer_clear_button.clicked.connect(self.new_transfer_form)
        self.transfers_tab.transfers_table.clicked.connect(self.on_transfer_selection_changed)
        self.transfers_tab.transfers_table.selectionModel().selectionChanged.connect(
            lambda *_: self.on_transfer_selection_changed()
        )
        self.transfers_tab.month_filter.currentTextChanged.connect(self.on_transfers_month_changed)
        self.accounts_tab.month_filter.currentTextChanged.connect(self.on_accounts_month_changed)
        self.accounts_tab.account_tabs.currentChanged.connect(self.on_accounts_subtab_changed)

        self.budget_tab.save_button.clicked.connect(self.save_budget_allocation)
        self.budget_tab.delete_button.clicked.connect(self.delete_budget_allocation)
        self.budget_tab.definitions_button.clicked.connect(self.show_budget_definitions_dialog)
        self.budget_tab.refresh_button.clicked.connect(self.refresh_budget_for_selected_month)
        self.budget_tab.table.clicked.connect(self.on_budget_selection_changed)
        self.budget_tab.table.selectionModel().selectionChanged.connect(
            lambda *_: self.on_budget_selection_changed()
        )
        self.budget_tab.month_picker.currentTextChanged.connect(self.on_budget_month_changed)

        self.bills_tab.save_button.clicked.connect(self.save_bill)
        self.bills_tab.delete_button.clicked.connect(self.delete_bill)
        self.bills_tab.refresh_bills_button.clicked.connect(self.refresh_bills_for_selected_month)
        self.bills_tab.bill_definitions_button.clicked.connect(self.show_bill_definitions_dialog)
        self.bills_tab.sort_name_button.clicked.connect(lambda: self.set_bills_sort("name"))
        self.bills_tab.sort_category_button.clicked.connect(lambda: self.set_bills_sort("category"))
        self.bills_tab.sort_due_button.clicked.connect(lambda: self.set_bills_sort("payment_due"))
        self.bills_tab.table.clicked.connect(self.on_bill_selection_changed)
        self.bills_tab.table.selectionModel().selectionChanged.connect(
            lambda *_: self.on_bill_selection_changed()
        )
        self.bills_tab.month_filter.currentTextChanged.connect(self.on_bills_month_changed)
        self.income_tab.save_button.clicked.connect(self.save_income)
        self.income_tab.delete_button.clicked.connect(self.delete_income)
        self.income_tab.refresh_income_button.clicked.connect(self.refresh_income_for_selected_month)
        self.income_tab.income_definitions_button.clicked.connect(self.show_income_definitions_dialog)
        self.income_tab.sort_category_button.clicked.connect(lambda: self.set_income_sort("category"))
        self.income_tab.sort_description_button.clicked.connect(lambda: self.set_income_sort("description"))
        self.income_tab.sort_account_button.clicked.connect(lambda: self.set_income_sort("account"))
        self.income_tab.table.clicked.connect(self.on_income_selection_changed)
        self.income_tab.table.selectionModel().selectionChanged.connect(
            lambda *_: self.on_income_selection_changed()
        )
        self.income_tab.month_filter.currentTextChanged.connect(self.on_income_month_changed)
        self.dashboard_tab.month_picker.currentTextChanged.connect(self.on_dashboard_month_changed)
        self.reports_tab.run_button.clicked.connect(self.run_selected_report)
        self.tabs.currentChanged.connect(self._on_tabs_current_changed)
        self.help_button.clicked.connect(self.show_help)
        self.settings_button.clicked.connect(self.show_settings_dialog)
        self.about_button.clicked.connect(self.show_about_dialog)
        self.exit_button.clicked.connect(self.close)

    def _on_tabs_current_changed(self, index: int) -> None:
        self._update_current_tab_header_label(index)
        if index >= 0 and self.tabs.widget(index) is self.dashboard_tab:
            self.refresh_dashboard()

    def _update_current_tab_header_label(self, index: int | None = None) -> None:
        if index is None:
            index = self.tabs.currentIndex()
        if index < 0:
            self.current_tab_header_label.setText("")
            return
        self.current_tab_header_label.setText(self.tabs.tabText(index))

    def import_transactions_xlsx(self) -> None:
        start_dir = self._import_dialog_start_dir()
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Import Transactions XLSX",
            start_dir,
            "Excel Files (*.xlsx *.xlsm *.xltx *.xltm);;All Files (*)",
        )
        if not file_path:
            return
        self._persist_last_import_dir(Path(file_path).parent)

        importer = XLSXTransactionImporter(
            self.context.transactions_service,
            self.context.categories_repo,
            self.context.accounts_repo,
            transfer_rules=self.context.settings.get("transfers", {}).get("rules", []),
            logger=self.logger,
        )

        try:
            result = importer.import_file(Path(file_path))
        except ValueError as exc:
            self.logger.error("XLSX import validation failed: %s", exc)
            QMessageBox.warning(self, "XLSX Import Error", str(exc))
            return
        except Exception as exc:
            self.logger.error("XLSX import failed: %s", exc)
            QMessageBox.critical(self, "XLSX Import Failed", str(exc))
            return

        month_text = result.import_period_key
        self._refresh_dashboard_month_filter(preferred_month=month_text)
        self._refresh_transactions_month_filter(preferred_month=month_text)
        self._refresh_transfers_month_filter(preferred_month=month_text)
        self._refresh_accounts_month_filter(preferred_month=month_text)
        self._refresh_budget_month_filter(preferred_month=month_text)
        self._refresh_bills_month_filter(preferred_month=month_text)
        self._refresh_transaction_form_choices()
        self.budget_tab.month_picker.blockSignals(True)
        self.budget_tab.month_picker.setCurrentText(month_text)
        self.budget_tab.month_picker.blockSignals(False)
        self.dashboard_tab.month_picker.blockSignals(True)
        self.dashboard_tab.month_picker.setCurrentText(month_text)
        self.dashboard_tab.month_picker.blockSignals(False)
        self._sync_reports_period_from_dashboard(preferred_month=month_text)
        self.refresh_transactions()
        self.refresh_transfers()
        self.refresh_accounts()
        self.refresh_budget_allocations()
        self.refresh_bills()
        self.refresh_income()
        self.refresh_dashboard()

        if result.transfer_rule_override_count > 0:
            examples = "\n".join(f"- {line}" for line in result.transfer_rule_override_examples)
            details = f"\n\nExamples:\n{examples}" if examples else ""
            QMessageBox.information(
                self,
                "Transfer Rule Override Applied",
                "One or more transfer rules overrode spreadsheet source accounts.\n"
                "Rule definitions take priority when a transfer rule matches.\n\n"
                f"Overrides applied: {result.transfer_rule_override_count}\n"
                "To fix an override, edit the transfer rule in Settings.\n"
                "Reminder: transfer rows are read-only on Transactions. "
                "Use the Transfers tab for manual transfer edits."
                f"{details}",
            )

        try:
            import_year, import_month = self._selected_year_month(self.dashboard_tab.month_picker)
        except Exception:  # noqa: BLE001
            import_year, import_month = self.transactions_view_year, self.transactions_view_month
        self._show_dashboard_category_mismatch_dialog(import_year, import_month)

        filename = Path(file_path).name
        self.statusBar().showMessage(
            f"Imported {result.imported_count} transactions from {filename} "
            f"(replaced {result.deleted_count} prior rows for {month_text})",
            7000,
        )
        self.logger.info(
            "Imported %s transactions from %s, replaced %s rows for month: %s",
            result.imported_count,
            file_path,
            result.deleted_count,
            month_text,
        )

    def _show_dashboard_category_mismatch_dialog(self, year: int, month: int) -> None:
        snapshot = self._compute_dashboard_snapshot_for_month(int(year), int(month))
        details = list(snapshot.get("relabeled_uncategorized_txn_details") or [])
        if not details:
            return

        lines = []
        for detail in details:
            lines.append(
                f"- txn_id={detail.get('txn_id')} | {detail.get('txn_date')} | "
                f"{detail.get('txn_type')} | {detail.get('txn_description')} | "
                f"category='{detail.get('category_name')}' (id={detail.get('category_id')}) | "
                f"expected={detail.get('expected_category_type')} | "
                f"reason={detail.get('mismatch_reason')}"
            )
        message = (
            "BudgetPal relabeled one or more transactions to 'Uncategorized' because the "
            "category type does not match transaction type.\n\n"
            "To fix this, update the category type (expense/income) in Settings > Definitions "
            "or correct the source transaction category.\n\n"
            f"Month: {int(year):04d}-{int(month):02d}\n"
            f"Mismatches: {len(details)}\n\n"
            "Details:\n"
            + "\n".join(lines)
        )
        QMessageBox.warning(self, "Category Type Mismatches", message)

    def _import_dialog_start_dir(self) -> str:
        ui_settings = self.context.settings.get("ui", {})
        raw = str(ui_settings.get("last_import_dir", "")).strip()
        if raw:
            candidate = Path(raw).expanduser()
            if candidate.is_file():
                candidate = candidate.parent
            if candidate.exists():
                return str(candidate)
        return str(Path.home())

    def _persist_last_import_dir(self, directory: Path) -> None:
        try:
            resolved = directory.expanduser().resolve()
        except OSError:
            resolved = directory

        ui_settings = self.context.settings.setdefault("ui", {})
        new_value = str(resolved)
        if str(ui_settings.get("last_import_dir", "")).strip() == new_value:
            return

        ui_settings["last_import_dir"] = new_value
        try:
            get_settings_manager().save(self.context.settings)
        except OSError as exc:
            self.logger.error("Failed to save last import directory: %s", exc)

    def _selected_year_month(self, selector) -> tuple[int, int]:
        raw = selector.currentText().strip()
        year_str, month_str = raw.split("-")
        return int(year_str), int(month_str)

    def refresh_transactions(self) -> None:
        rows = self.context.transactions_service.list_for_month(
            year=self.transactions_view_year,
            month=self.transactions_view_month,
            limit=5000,
        )

        expense_rows: list[dict] = []
        income_rows: list[dict] = []
        transfer_count = 0
        for row in rows:
            data = dict(row)
            amount_cents = int(data.get("amount_cents") or 0)
            txn_type = str(data.get("txn_type") or "").strip().lower()
            if txn_type == "transfer":
                transfer_count += 1
            data["display_amount_cents"] = abs(amount_cents)
            data["description_display"] = self._normalized_display_text(data.get("description"))

            # Pure transfer semantics on Transactions tab:
            # - show transfer outflow leg on Expenses side only
            # - do not show transfer inflow leg on Income side
            if txn_type == "income":
                income_rows.append(data)
            elif txn_type == "expense":
                expense_rows.append(data)
            elif txn_type == "transfer":
                if amount_cents < 0:
                    expense_rows.append(data)
            elif amount_cents > 0:
                income_rows.append(data)
            elif amount_cents < 0:
                expense_rows.append(data)

        expense_rows.sort(
            key=lambda r: (str(r.get("txn_date", "")), int(r.get("txn_id") or 0)),
        )
        income_rows.sort(
            key=lambda r: (str(r.get("txn_date", "")), int(r.get("txn_id") or 0)),
        )
        if expense_rows:
            expense_rows.append(self._transactions_spacer_row())
        if income_rows:
            income_rows.append(self._transactions_spacer_row())

        self.transactions_tab.expense_model.replace_rows(expense_rows)
        self.transactions_tab.income_model.replace_rows(income_rows)
        self.transactions_tab.ensure_bottom_rows_visible()
        self.transactions_tab.updateGeometry()
        self.logger.info(
            "Loaded %s transactions for %s-%02d (%s expenses, %s income, %s transfers)",
            len(rows),
            self.transactions_view_year,
            self.transactions_view_month,
            len(expense_rows),
            len(income_rows),
            transfer_count,
        )
        self.refresh_dashboard()

    def refresh_transfers(self) -> None:
        rows = self.context.transactions_service.list_transfer_summaries_for_month(
            year=self.transfers_view_year,
            month=self.transfers_view_month,
            limit=2000,
        )
        self.transfers_tab.transfers_model.replace_rows([dict(row) for row in rows])
        self.logger.info(
            "Loaded %s transfers for %s-%02d.",
            len(rows),
            self.transfers_view_year,
            self.transfers_view_month,
        )

    def _connect_account_pane_signals(self, pane) -> None:
        pane.beginning_balance_save_requested.connect(
            self.on_account_beginning_balance_save_requested
        )
        pane.statement_save_requested.connect(self.on_account_statement_save_requested)
        pane.txn_cleared_toggled.connect(self.on_account_txn_cleared_toggled)
        pane.txn_note_edited.connect(self.on_account_txn_note_edited)
        pane.sort_changed.connect(self.on_account_sort_changed)
        pane.select_all_requested.connect(self.on_account_select_all_requested)
        pane.clear_all_requested.connect(self.on_account_clear_all_requested)

    def refresh_accounts(self) -> None:
        account_rows = self.context.accounts_repo.list_active(
            include_external=True,
            include_hidden_from_accounts_tab=False,
        )
        created_panes = self.accounts_tab.sync_accounts(account_rows)
        for pane in created_panes:
            self._connect_account_pane_signals(pane)

        self._set_accounts_view_month(self.accounts_view_year, self.accounts_view_month)

        for row in account_rows:
            account_id = int(row["account_id"])
            account_type = str(row.get("account_type") or "").strip().lower()
            pane = self.accounts_tab.pane_for_account_id(account_id)
            if pane is None:
                continue

            beginning_balance_cents = (
                self.context.transactions_service.get_account_month_beginning_balance(
                    year=self.accounts_view_year,
                    month=self.accounts_view_month,
                    account_id=account_id,
                )
            )
            pane.beginning_balance_input.blockSignals(True)
            pane.beginning_balance_input.setText(f"{beginning_balance_cents / 100:.2f}")
            pane.beginning_balance_input.blockSignals(False)

            include_prior_uncleared = account_type in {"checking", "credit"}
            ledger_rows = self.context.transactions_service.list_account_ledger_for_month(
                year=self.accounts_view_year,
                month=self.accounts_view_month,
                account_id=account_id,
                include_prior_uncleared=include_prior_uncleared,
                limit=10000,
            )

            sort_key = pane.sort_key()
            if sort_key == "type":
                ledger_rows.sort(
                    key=lambda item: (
                        str(item.get("payment_type") or item.get("txn_type") or "")
                        .strip()
                        .casefold(),
                        str(item.get("txn_date") or ""),
                        int(item.get("txn_id") or 0),
                    )
                )
            elif sort_key == "txn_id":
                ledger_rows.sort(
                    key=lambda item: int(item.get("txn_id") or 0)
                )
            else:
                ledger_rows.sort(
                    key=lambda item: (
                        str(item.get("txn_date") or ""),
                        int(item.get("txn_id") or 0),
                    )
                )

            running_balance_cents = int(beginning_balance_cents)
            table_rows: list[dict] = []
            for ledger_row in ledger_rows:
                data = dict(ledger_row)
                txn_type = str(data.get("txn_type") or "").strip().lower()
                raw_amount_cents = int(data.get("amount_cents") or 0)
                if txn_type == "income":
                    signed_amount_cents = abs(raw_amount_cents)
                elif txn_type == "expense":
                    signed_amount_cents = -abs(raw_amount_cents)
                else:
                    signed_amount_cents = raw_amount_cents
                running_balance_cents += signed_amount_cents
                data["payment_type_display"] = self._normalized_display_text(data.get("payment_type"))
                data["description_display"] = self._normalized_display_text(data.get("description"))
                data["note_display"] = self._normalized_display_text(data.get("note"))
                data["amount_display"] = self._format_currency_signed(signed_amount_cents)
                data["running_balance_display"] = self._format_currency_balance(running_balance_cents)
                table_rows.append(data)

            pane.model.replace_rows(table_rows)
            pane.ending_balance_value.setText(self._format_currency_balance(running_balance_cents))

            statement_row = self.context.transactions_service.get_account_month_statement(
                year=self.accounts_view_year,
                month=self.accounts_view_month,
                account_id=account_id,
            )
            statement_ending_cents = statement_row.get("statement_ending_balance_cents")
            statement_date_text = str(statement_row.get("statement_ending_date") or "").strip()
            reported_current_cents = statement_row.get("reported_current_balance_cents")
            reported_available_cents = statement_row.get("reported_available_credit_cents")
            line_of_credit_cents = row.get("line_of_credit_cents")
            statement_ending_display = (
                self._format_currency_balance(int(statement_ending_cents))
                if statement_ending_cents is not None
                else ""
            )
            pane.set_statement_fields(
                statement_ending_display,
                statement_date_text,
                line_of_credit_text=(
                    self._format_currency_balance(int(line_of_credit_cents))
                    if line_of_credit_cents is not None
                    else "$0.00"
                ),
                reported_current_balance_text=(
                    self._format_currency_balance(int(reported_current_cents))
                    if reported_current_cents is not None
                    else ""
                ),
                reported_available_credit_text=(
                    self._format_currency_balance(int(reported_available_cents))
                    if reported_available_cents is not None
                    else ""
                ),
            )

            pending_deposits_cents = 0
            pending_withdrawals_cents = 0
            for table_row in table_rows:
                if bool(table_row.get("is_cleared")):
                    continue
                signed_amount_cents = int(table_row.get("amount_cents") or 0)
                if signed_amount_cents > 0:
                    pending_deposits_cents += signed_amount_cents
                elif signed_amount_cents < 0:
                    pending_withdrawals_cents += abs(signed_amount_cents)
            net_pending_cents = pending_deposits_cents - pending_withdrawals_cents
            cleared_register_cents = running_balance_cents - net_pending_cents

            adjusted_statement_display = "N/A"
            difference_display = "N/A"
            status_text = "Statement balance not entered"
            status_ok: bool | None = None
            if statement_ending_cents is not None:
                adjusted_statement_cents = int(statement_ending_cents) + net_pending_cents
                difference_cents = adjusted_statement_cents - running_balance_cents
                adjusted_statement_display = self._format_currency_balance(adjusted_statement_cents)
                difference_display = self._format_currency_signed(difference_cents)
                if difference_cents == 0:
                    status_text = "Balanced"
                    status_ok = True
                else:
                    status_text = "Out of balance"
                    status_ok = False

            if account_type == "credit":
                computed_current_debt_cents: int | None = None
                computed_available_credit_cents: int | None = None
                diff_cents: int | None = None
                credit_status = "Statement balance not entered"
                credit_status_ok: bool | None = None
                if statement_ending_cents is not None:
                    statement_date_value = None
                    if statement_date_text:
                        try:
                            statement_date_value = datetime.strptime(
                                statement_date_text,
                                "%Y-%m-%d",
                            ).date()
                        except ValueError:
                            statement_date_value = None

                    debt_delta_cents = 0
                    for table_row in table_rows:
                        if not bool(table_row.get("is_cleared")):
                            continue
                        txn_date_raw = str(table_row.get("txn_date") or "").strip()
                        if statement_date_value is not None and txn_date_raw:
                            try:
                                txn_date_value = datetime.strptime(txn_date_raw, "%Y-%m-%d").date()
                            except ValueError:
                                continue
                            if txn_date_value <= statement_date_value:
                                continue
                        signed_amount_cents = int(table_row.get("amount_cents") or 0)
                        debt_delta_cents += -signed_amount_cents

                    computed_current_debt_cents = int(statement_ending_cents) + debt_delta_cents
                    if line_of_credit_cents is not None:
                        computed_available_credit_cents = int(line_of_credit_cents) - int(
                            computed_current_debt_cents
                        )

                    if reported_available_cents is not None and computed_available_credit_cents is not None:
                        diff_cents = int(reported_available_cents) - int(computed_available_credit_cents)
                        if diff_cents == 0:
                            credit_status = "Reconciled"
                            credit_status_ok = True
                        else:
                            credit_status = "Needs Review"
                            credit_status_ok = False
                    elif line_of_credit_cents is None:
                        credit_status = "Line of credit not set"
                        credit_status_ok = None
                    else:
                        credit_status = "Reported available credit not entered"
                        credit_status_ok = None

                    if (
                        line_of_credit_cents is not None
                        and reported_available_cents is not None
                        and int(reported_available_cents) == int(line_of_credit_cents)
                    ):
                        credit_status = (
                            f"{credit_status} | Paid in Full"
                            if credit_status
                            else "Paid in Full"
                        )

                pane.set_credit_reconciliation_values(
                    computed_current_balance_display=(
                        self._format_currency_balance(int(computed_current_debt_cents))
                        if computed_current_debt_cents is not None
                        else "N/A"
                    ),
                    computed_available_credit_display=(
                        self._format_currency_balance(int(computed_available_credit_cents))
                        if computed_available_credit_cents is not None
                        else "N/A"
                    ),
                    difference_display=(
                        self._format_currency_signed(int(diff_cents))
                        if diff_cents is not None
                        else "N/A"
                    ),
                    status_text=credit_status,
                    status_ok=credit_status_ok,
                )
            else:
                pane.set_reconciliation_values(
                    pending_deposits_display=self._format_currency(pending_deposits_cents),
                    pending_withdrawals_display=self._format_currency(pending_withdrawals_cents),
                    net_pending_display=self._format_currency_signed(net_pending_cents),
                    cleared_register_display=self._format_currency_balance(cleared_register_cents),
                    adjusted_statement_display=adjusted_statement_display,
                    difference_display=difference_display,
                    status_text=status_text,
                    status_ok=status_ok,
                )

        self._refresh_accounts_detail_panel()

        self.logger.info(
            "Refreshed Accounts tab for %s-%02d (%s account tabs).",
            self.accounts_view_year,
            self.accounts_view_month,
            len(account_rows),
        )

    def _refresh_accounts_detail_panel(self) -> None:
        pane = self.accounts_tab.current_pane()
        if pane is None:
            self.accounts_tab.set_account_details(None)
            return
        row = self.accounts_tab.account_row_by_id(int(pane.account_id))
        self.accounts_tab.set_account_details(row)

    def on_account_beginning_balance_save_requested(self, account_id: int, amount_text: str) -> None:
        try:
            beginning_balance_cents = self._parse_currency_cents_allow_negative(amount_text)
        except ValueError as exc:
            QMessageBox.warning(self, "Beginning Balance", str(exc))
            self.refresh_accounts()
            return

        existing_cents = self.context.transactions_service.get_account_month_beginning_balance(
            year=self.accounts_view_year,
            month=self.accounts_view_month,
            account_id=int(account_id),
        )
        if int(existing_cents) == int(beginning_balance_cents):
            return

        self.context.transactions_service.set_account_month_beginning_balance(
            year=self.accounts_view_year,
            month=self.accounts_view_month,
            beginning_balance_cents=beginning_balance_cents,
            account_id=int(account_id),
        )
        self.statusBar().showMessage(
            f"Saved account beginning balance for {self.accounts_view_year:04d}-{self.accounts_view_month:02d}.",
            3000,
        )
        self.refresh_accounts()

    def on_account_statement_save_requested(
        self,
        account_id: int,
        statement_balance_text: str,
        statement_date_text: str,
        reported_current_balance_text: str,
        reported_available_credit_text: str,
    ) -> None:
        balance_text = str(statement_balance_text or "").strip()
        date_text = str(statement_date_text or "").strip()
        reported_current_text = str(reported_current_balance_text or "").strip()
        reported_available_text = str(reported_available_credit_text or "").strip()

        account_row = self.accounts_tab.account_row_by_id(int(account_id)) or {}
        account_type = str(account_row.get("account_type") or "").strip().lower()

        statement_ending_balance_cents: int | None
        if not balance_text:
            statement_ending_balance_cents = None
        else:
            try:
                statement_ending_balance_cents = self._parse_currency_cents_allow_negative(balance_text)
            except ValueError as exc:
                QMessageBox.warning(self, "Statement Ending Balance", str(exc))
                self.refresh_accounts()
                return
            if account_type == "credit":
                statement_ending_balance_cents = abs(int(statement_ending_balance_cents))

        reported_current_balance_cents: int | None = None
        reported_available_credit_cents: int | None = None
        if account_type == "credit":
            if reported_current_text:
                try:
                    reported_current_balance_cents = abs(
                        int(self._parse_currency_cents_allow_negative(reported_current_text))
                    )
                except ValueError as exc:
                    QMessageBox.warning(self, "Reported Current Balance", str(exc))
                    self.refresh_accounts()
                    return
            if reported_available_text:
                try:
                    reported_available_credit_cents = abs(
                        int(self._parse_currency_cents_allow_negative(reported_available_text))
                    )
                except ValueError as exc:
                    QMessageBox.warning(self, "Reported Available Credit", str(exc))
                    self.refresh_accounts()
                    return

        if date_text:
            try:
                datetime.strptime(date_text, "%Y-%m-%d")
            except ValueError:
                QMessageBox.warning(
                    self,
                    "Statement Date",
                    "Statement Date must be in YYYY-MM-DD format.",
                )
                self.refresh_accounts()
                return
        else:
            date_text = ""

        self.context.transactions_service.set_account_month_statement(
            year=self.accounts_view_year,
            month=self.accounts_view_month,
            account_id=int(account_id),
            statement_ending_balance_cents=statement_ending_balance_cents,
            statement_ending_date=(date_text or None),
            reported_current_balance_cents=reported_current_balance_cents,
            reported_available_credit_cents=reported_available_credit_cents,
        )
        self.logger.info(
            "Saved statement reconciliation values for account_id=%s, month=%s-%02d.",
            account_id,
            self.accounts_view_year,
            self.accounts_view_month,
        )
        self.statusBar().showMessage("Saved statement reconciliation values.", 3000)
        self.refresh_accounts()

    def on_account_txn_cleared_toggled(self, txn_id: int, is_cleared: bool) -> None:
        self._set_transaction_cleared(txn_id, is_cleared)

    def on_account_txn_note_edited(self, txn_id: int, note_text: str) -> None:
        updated = self.context.transactions_service.set_transaction_note(
            txn_id=int(txn_id),
            note=note_text,
        )
        if not updated:
            QMessageBox.warning(
                self,
                "Update Note",
                "The selected transaction no longer exists.",
            )
            self.refresh_accounts()
            return
        self.logger.info("Updated transaction %s note.", txn_id)
        self.statusBar().showMessage("Updated transaction note.", 2500)
        self.refresh_transactions()
        self.refresh_transfers()
        self.refresh_accounts()

    def on_account_sort_changed(self, account_id: int, _sort_key: str) -> None:
        current = self.accounts_tab.current_pane()
        if current is not None and int(current.account_id) == int(account_id):
            self.refresh_accounts()

    def _set_all_account_rows(self, account_id: int, is_cleared: bool) -> None:
        pane = self.accounts_tab.pane_for_account_id(int(account_id))
        if pane is None:
            return
        changed = 0
        for index in range(pane.model.rowCount()):
            row = pane.model.row_dict(index) or {}
            txn_id = int(row.get("txn_id") or 0)
            if txn_id <= 0:
                continue
            current = bool(row.get("is_cleared"))
            if current == bool(is_cleared):
                continue
            updated = self.context.transactions_service.set_transaction_cleared(
                txn_id=txn_id,
                is_cleared=bool(is_cleared),
            )
            if updated:
                changed += 1
        if changed:
            self.refresh_accounts()
        self.statusBar().showMessage(f"Updated {changed} account rows.", 3000)

    def on_account_select_all_requested(self, account_id: int) -> None:
        self._set_all_account_rows(account_id, True)

    def on_account_clear_all_requested(self, account_id: int) -> None:
        self._set_all_account_rows(account_id, False)

    def _set_transaction_cleared(self, txn_id: int, is_cleared: bool) -> None:
        updated = self.context.transactions_service.set_transaction_cleared(
            txn_id=int(txn_id),
            is_cleared=bool(is_cleared),
        )
        if not updated:
            QMessageBox.warning(
                self,
                "Update Cleared Status",
                "The selected transaction no longer exists.",
            )
            self.refresh_accounts()
            return
        self.logger.info("Updated transaction %s cleared status to %s", txn_id, int(is_cleared))
        self.statusBar().showMessage("Updated cleared status.", 3000)
        self.refresh_accounts()

    def _set_transactions_view_month(self, year: int, month: int) -> None:
        self.transactions_view_year = year
        self.transactions_view_month = month
        self.transactions_tab.view_heading.setText(
            f"Transactions for {self.transactions_view_year}-{self.transactions_view_month:02d}"
        )

    def _set_transfers_view_month(self, year: int, month: int) -> None:
        self.transfers_view_year = year
        self.transfers_view_month = month
        self.transfers_tab.view_heading.setText(
            f"Transfers for {self.transfers_view_year}-{self.transfers_view_month:02d}"
        )

    def _set_budget_view_month(self, year: int, month: int) -> None:
        self.budget_view_year = year
        self.budget_view_month = month
        self.budget_tab.view_heading.setText(
            f"Budget Allocations for {self.budget_view_year}-{self.budget_view_month:02d}"
        )

    def _set_bills_view_month(self, year: int, month: int) -> None:
        self.bills_view_year = year
        self.bills_view_month = month
        self.bills_tab.view_heading.setText(
            f"Bills for {self.bills_view_year}-{self.bills_view_month:02d}"
        )

    def _set_income_view_month(self, year: int, month: int) -> None:
        self.income_view_year = year
        self.income_view_month = month
        self.income_tab.view_heading.setText(
            f"Income for {self.income_view_year}-{self.income_view_month:02d}"
        )

    def _set_accounts_view_month(self, year: int, month: int) -> None:
        self.accounts_view_year = year
        self.accounts_view_month = month
        pane = self.accounts_tab.current_pane()
        account_label = str(getattr(pane, "account_name", "") or "").strip()
        suffix = f" • {account_label}" if account_label else ""
        self.accounts_tab.view_heading.setText(
            f"Accounts for {self.accounts_view_year}-{self.accounts_view_month:02d}{suffix}"
        )

    @staticmethod
    def _month_key(year: int, month: int) -> str:
        return f"{int(year):04d}-{int(month):02d}"

    @staticmethod
    def _normalized_source_system(source_system: str | None) -> str:
        source = str(source_system or "").strip().lower()
        if source == "subtracker":
            return "subtracker"
        return "budgetpal"

    def _mark_bills_month_dirty(self, year: int, month: int, source_system: str | None) -> None:
        key = self._month_key(year, month)
        source = self._normalized_source_system(source_system)
        if source == "subtracker":
            self._subscriptions_dirty_by_month[key] = True
        else:
            self._bills_dirty_by_month[key] = True

    def _is_bills_month_dirty(self, year: int, month: int) -> bool:
        key = self._month_key(year, month)
        return bool(self._bills_dirty_by_month.get(key) or self._subscriptions_dirty_by_month.get(key))

    def _is_subscriptions_month_dirty(self, year: int, month: int) -> bool:
        key = self._month_key(year, month)
        return bool(self._subscriptions_dirty_by_month.get(key))

    def _clear_bills_month_dirty(self, year: int, month: int, source_system: str | None = None) -> None:
        key = self._month_key(year, month)
        if source_system is None:
            self._bills_dirty_by_month[key] = False
            self._subscriptions_dirty_by_month[key] = False
            return
        source = self._normalized_source_system(source_system)
        if source == "subtracker":
            self._subscriptions_dirty_by_month[key] = False
        else:
            self._bills_dirty_by_month[key] = False

    def _mark_income_month_dirty(self, year: int, month: int) -> None:
        key = self._month_key(year, month)
        self._income_dirty_by_month[key] = True

    def _is_income_month_dirty(self, year: int, month: int) -> bool:
        key = self._month_key(year, month)
        return bool(self._income_dirty_by_month.get(key))

    def _clear_income_month_dirty(self, year: int, month: int) -> None:
        key = self._month_key(year, month)
        self._income_dirty_by_month[key] = False

    def _mark_budget_month_dirty(self, year: int, month: int) -> None:
        key = self._month_key(year, month)
        self._budget_dirty_by_month[key] = True

    def _is_budget_month_dirty(self, year: int, month: int) -> bool:
        key = self._month_key(year, month)
        return bool(self._budget_dirty_by_month.get(key))

    def _clear_budget_month_dirty(self, year: int, month: int) -> None:
        key = self._month_key(year, month)
        self._budget_dirty_by_month[key] = False

    def _refresh_transactions_month_filter(self, preferred_month: str | None = None) -> None:
        current_month = self.transactions_tab.month_filter.currentText().strip()
        months = self.context.transactions_service.list_available_months()
        default_month = date.today().strftime("%Y-%m")

        month_set = set(months)
        month_set.update(self._rolling_month_labels(months_back=12, months_forward=12))
        month_set.add(default_month)
        if current_month:
            month_set.add(current_month)
        if preferred_month:
            month_set.add(preferred_month)
        month_values = sorted(month_set, reverse=True)

        self.transactions_tab.month_filter.blockSignals(True)
        self.transactions_tab.month_filter.clear()
        self.transactions_tab.month_filter.addItems(month_values)

        target_month = preferred_month or current_month or default_month
        if target_month not in month_set:
            target_month = month_values[0]
        self.transactions_tab.month_filter.setCurrentText(target_month)
        self.transactions_tab.month_filter.blockSignals(False)

        year_str, month_str = target_month.split("-")
        self._set_transactions_view_month(int(year_str), int(month_str))

    def _refresh_transfers_month_filter(self, preferred_month: str | None = None) -> None:
        current_month = self.transfers_tab.month_filter.currentText().strip()
        months = self.context.transactions_service.list_available_months()
        default_month = date.today().strftime("%Y-%m")

        month_set = set(months)
        month_set.update(self._rolling_month_labels(months_back=12, months_forward=12))
        month_set.add(default_month)
        if current_month:
            month_set.add(current_month)
        if preferred_month:
            month_set.add(preferred_month)
        month_values = sorted(month_set, reverse=True)

        self.transfers_tab.month_filter.blockSignals(True)
        self.transfers_tab.month_filter.clear()
        self.transfers_tab.month_filter.addItems(month_values)

        target_month = preferred_month or current_month or default_month
        if target_month not in month_set:
            target_month = month_values[0]
        self.transfers_tab.month_filter.setCurrentText(target_month)
        self.transfers_tab.month_filter.blockSignals(False)

        year_str, month_str = target_month.split("-")
        self._set_transfers_view_month(int(year_str), int(month_str))

    def _refresh_budget_month_filter(self, preferred_month: str | None = None) -> None:
        current_month = self.budget_tab.month_picker.currentText().strip()
        months = self.context.transactions_service.list_available_months()
        months.extend(self.context.budget_allocations_service.list_available_months())
        default_month = date.today().strftime("%Y-%m")

        month_set = set(months)
        month_set.update(self._rolling_month_labels(months_back=12, months_forward=12))
        month_set.add(default_month)
        if current_month:
            month_set.add(current_month)
        if preferred_month:
            month_set.add(preferred_month)
        month_values = sorted(month_set, reverse=True)

        self.budget_tab.month_picker.blockSignals(True)
        self.budget_tab.month_picker.clear()
        self.budget_tab.month_picker.addItems(month_values)

        target_month = preferred_month or current_month or default_month
        if target_month not in month_set:
            target_month = month_values[0]
        self.budget_tab.month_picker.setCurrentText(target_month)
        self.budget_tab.month_picker.blockSignals(False)

        year_str, month_str = target_month.split("-")
        self._set_budget_view_month(int(year_str), int(month_str))

    def _refresh_dashboard_month_filter(self, preferred_month: str | None = None) -> None:
        current_month = self.dashboard_tab.month_picker.currentText().strip()
        months = self.context.transactions_service.list_available_months()
        default_month = date.today().strftime("%Y-%m")

        month_set = set(months)
        month_set.update(self._rolling_month_labels(months_back=12, months_forward=12))
        month_set.add(default_month)
        if current_month:
            month_set.add(current_month)
        if preferred_month:
            month_set.add(preferred_month)
        month_values = sorted(month_set, reverse=True)

        self.dashboard_tab.month_picker.blockSignals(True)
        self.dashboard_tab.month_picker.clear()
        self.dashboard_tab.month_picker.addItems(month_values)

        target_month = preferred_month or current_month or default_month
        if target_month not in month_set:
            target_month = month_values[0]
        self.dashboard_tab.month_picker.setCurrentText(target_month)
        self.dashboard_tab.month_picker.blockSignals(False)

    def _refresh_bills_month_filter(self, preferred_month: str | None = None) -> None:
        current_month = self.bills_tab.month_filter.currentText().strip()
        months = self.context.transactions_service.list_available_months()
        default_month = date.today().strftime("%Y-%m")

        month_set = set(months)
        month_set.update(self._rolling_month_labels(months_back=12, months_forward=12))
        month_set.add(default_month)
        if current_month:
            month_set.add(current_month)
        if preferred_month:
            month_set.add(preferred_month)
        month_values = sorted(month_set, reverse=True)

        self.bills_tab.month_filter.blockSignals(True)
        self.bills_tab.month_filter.clear()
        self.bills_tab.month_filter.addItems(month_values)

        target_month = preferred_month or current_month or default_month
        if target_month not in month_set:
            target_month = month_values[0]
        self.bills_tab.month_filter.setCurrentText(target_month)
        self.bills_tab.month_filter.blockSignals(False)

        year_str, month_str = target_month.split("-")
        self._set_bills_view_month(int(year_str), int(month_str))

    def _refresh_income_month_filter(self, preferred_month: str | None = None) -> None:
        current_month = self.income_tab.month_filter.currentText().strip()
        months = self.context.transactions_service.list_available_months()
        default_month = date.today().strftime("%Y-%m")

        month_set = set(months)
        month_set.update(self._rolling_month_labels(months_back=12, months_forward=12))
        month_set.add(default_month)
        if current_month:
            month_set.add(current_month)
        if preferred_month:
            month_set.add(preferred_month)
        month_values = sorted(month_set, reverse=True)

        self.income_tab.month_filter.blockSignals(True)
        self.income_tab.month_filter.clear()
        self.income_tab.month_filter.addItems(month_values)

        target_month = preferred_month or current_month or default_month
        if target_month not in month_set:
            target_month = month_values[0]
        self.income_tab.month_filter.setCurrentText(target_month)
        self.income_tab.month_filter.blockSignals(False)

        year_str, month_str = target_month.split("-")
        self._set_income_view_month(int(year_str), int(month_str))

    def _refresh_accounts_month_filter(self, preferred_month: str | None = None) -> None:
        current_month = self.accounts_tab.month_filter.currentText().strip()
        months = self.context.transactions_service.list_available_months()
        default_month = date.today().strftime("%Y-%m")

        month_set = set(months)
        month_set.update(self._rolling_month_labels(months_back=12, months_forward=12))
        month_set.add(default_month)
        if current_month:
            month_set.add(current_month)
        if preferred_month:
            month_set.add(preferred_month)
        month_values = sorted(month_set, reverse=True)

        self.accounts_tab.month_filter.blockSignals(True)
        self.accounts_tab.month_filter.clear()
        self.accounts_tab.month_filter.addItems(month_values)

        target_month = preferred_month or current_month or default_month
        if target_month not in month_set:
            target_month = month_values[0]
        self.accounts_tab.month_filter.setCurrentText(target_month)
        self.accounts_tab.month_filter.blockSignals(False)

        year_str, month_str = target_month.split("-")
        self._set_accounts_view_month(int(year_str), int(month_str))

    def _sync_reports_period_from_dashboard(self, preferred_month: str | None = None) -> None:
        month_value = (preferred_month or self.dashboard_tab.month_picker.currentText() or "").strip()
        if not month_value:
            return
        try:
            year_str, month_str = month_value.split("-")
        except ValueError:
            return
        self.reports_tab.year_picker.setCurrentText(str(int(year_str)))
        self.reports_tab.month_picker.setCurrentText(f"{int(month_str):02d}")

    def _load_static_reports(self) -> None:
        self.reports_tab.set_report_rows(report_rows())

    def on_transactions_month_changed(self, month_value: str) -> None:
        value = month_value.strip()
        if not value:
            return
        try:
            year_str, month_str = value.split("-")
            self._set_transactions_view_month(int(year_str), int(month_str))
        except ValueError:
            self.logger.warning("Invalid month filter value: %s", value)
            return
        self.refresh_transactions()

    def on_transfers_month_changed(self, month_value: str) -> None:
        value = month_value.strip()
        if not value:
            return
        try:
            year_str, month_str = value.split("-")
            self._set_transfers_view_month(int(year_str), int(month_str))
        except ValueError:
            self.logger.warning("Invalid transfers month filter value: %s", value)
            return
        self.refresh_transfers()

    def on_budget_month_changed(self, month_value: str) -> None:
        value = month_value.strip()
        if not value:
            return
        try:
            year_str, month_str = value.split("-")
            self._set_budget_view_month(int(year_str), int(month_str))
        except ValueError:
            self.logger.warning("Invalid budget month filter value: %s", value)
            return
        self.refresh_budget_allocations()

    def on_bills_month_changed(self, month_value: str) -> None:
        value = month_value.strip()
        if not value:
            return
        try:
            year_str, month_str = value.split("-")
            self._set_bills_view_month(int(year_str), int(month_str))
        except ValueError:
            self.logger.warning("Invalid bills month filter value: %s", value)
            return
        self.refresh_bills()

    def on_income_month_changed(self, month_value: str) -> None:
        value = month_value.strip()
        if not value:
            return
        try:
            year_str, month_str = value.split("-")
            self._set_income_view_month(int(year_str), int(month_str))
        except ValueError:
            self.logger.warning("Invalid income month filter value: %s", value)
            return
        self.refresh_income()

    def on_accounts_month_changed(self, month_value: str) -> None:
        value = month_value.strip()
        if not value:
            return
        try:
            year_str, month_str = value.split("-")
            self._set_accounts_view_month(int(year_str), int(month_str))
        except ValueError:
            self.logger.warning("Invalid accounts month filter value: %s", value)
            return
        self.refresh_accounts()

    def on_accounts_subtab_changed(self, _index: int) -> None:
        self._set_accounts_view_month(self.accounts_view_year, self.accounts_view_month)
        self._refresh_accounts_detail_panel()

    def on_dashboard_month_changed(self, month_value: str) -> None:
        value = month_value.strip()
        if not value:
            return
        # Keep Bills month in sync with Dashboard month selection for planning workflows.
        self._refresh_budget_month_filter(preferred_month=value)
        self._refresh_bills_month_filter(preferred_month=value)
        self._refresh_income_month_filter(preferred_month=value)
        self._refresh_accounts_month_filter(preferred_month=value)
        self._refresh_transactions_month_filter(preferred_month=value)
        self._refresh_transfers_month_filter(preferred_month=value)
        self._sync_reports_period_from_dashboard(preferred_month=value)
        self.refresh_transactions()
        self.refresh_transfers()
        self.refresh_budget_allocations()
        self.refresh_bills()
        self.refresh_income()
        self.refresh_accounts()
        self.refresh_dashboard()

    @staticmethod
    def _format_currency(cents: int) -> str:
        return f"${int(cents) / 100:,.2f}"

    @staticmethod
    def _format_currency_signed(cents: int) -> str:
        value = int(cents)
        if value > 0:
            return f"+${value / 100:,.2f}"
        if value < 0:
            return f"-${abs(value) / 100:,.2f}"
        return "$0.00"

    @staticmethod
    def _format_currency_balance(cents: int) -> str:
        value = int(cents)
        if value < 0:
            return f"-${abs(value) / 100:,.2f}"
        return f"${value / 100:,.2f}"

    @staticmethod
    def _parse_currency_cents_allow_negative(amount_text: str) -> int:
        cleaned = amount_text.strip().replace("$", "").replace(",", "")
        if not cleaned:
            return 0
        try:
            amount = Decimal(cleaned)
        except InvalidOperation as exc:
            raise ValueError("Starting Balance must be numeric (example: 5000.00).") from exc
        return int((amount * 100).quantize(Decimal("1"), rounding=ROUND_HALF_UP))

    @staticmethod
    def _dashboard_category_label(raw_value: object) -> str:
        text = str(raw_value or "").strip()
        return text or "Uncategorized"

    @classmethod
    def _build_dashboard_category_rows(
        cls,
        planned_by_category: dict[str, int],
        actual_by_category: dict[str, int],
        *,
        is_income: bool,
    ) -> list[dict]:
        rows: list[dict] = []
        planned_total = 0
        actual_total = 0
        categories = sorted(set(planned_by_category) | set(actual_by_category), key=str.casefold)
        for category_name in categories:
            planned_cents = int(planned_by_category.get(category_name, 0))
            actual_cents = int(actual_by_category.get(category_name, 0))
            if planned_cents == 0 and actual_cents == 0:
                continue
            planned_total += planned_cents
            actual_total += actual_cents
            diff_cents = (actual_cents - planned_cents) if is_income else (planned_cents - actual_cents)
            rows.append(
                {
                    "category_name": category_name,
                    "planned_display": cls._format_currency(planned_cents),
                    "actual_display": cls._format_currency(actual_cents),
                    "diff_display": cls._format_currency_signed(diff_cents),
                }
            )

        total_diff_cents = (actual_total - planned_total) if is_income else (planned_total - actual_total)
        total_row = {
            "category_name": "Totals",
            "planned_display": cls._format_currency(planned_total),
            "actual_display": cls._format_currency(actual_total),
            "diff_display": cls._format_currency_signed(total_diff_cents),
        }
        separator_row = {
            "category_name": "",
            "planned_display": "",
            "actual_display": "",
            "diff_display": "",
        }
        return [total_row, separator_row, *rows]

    def _compute_dashboard_snapshot_for_month(self, year: int, month: int) -> dict:
        month_row = self.context.budgeting_service.get_month(year, month)
        starting_balance_cents = int(month_row.get("starting_balance_cents") or 0)

        expense_category_rows = self.context.categories_repo.list_active(category_type="expense")
        income_category_rows = self.context.categories_repo.list_active(category_type="income")
        expense_category_ids = {int(row["category_id"]) for row in expense_category_rows}
        income_category_ids = {int(row["category_id"]) for row in income_category_rows}
        uncategorized_label = "Uncategorized"
        expense_category_names = {
            int(row["category_id"]): self._dashboard_category_label(row.get("name"))
            for row in expense_category_rows
        }
        income_category_names = {
            int(row["category_id"]): self._dashboard_category_label(row.get("name"))
            for row in income_category_rows
        }
        internal_account_ids = {
            int(row["account_id"])
            for row in self.context.accounts_repo.list_active(include_external=False)
        }

        planned_expense_by_category: dict[str, int] = {}
        for row in self.context.budget_allocations_service.list_month_allocations(year=year, month=month):
            category_id = row.get("category_id")
            cents = int(row.get("planned_cents") or 0)
            if cents == 0:
                continue
            if category_id is not None and int(category_id) in expense_category_ids:
                category_name = expense_category_names.get(
                    int(category_id),
                    self._dashboard_category_label(row.get("category_name")),
                )
            else:
                category_name = uncategorized_label
            planned_expense_by_category[category_name] = planned_expense_by_category.get(category_name, 0) + cents

        planned_income_by_category: dict[str, int] = {}
        for row in self.context.income_service.list_month_income(year=year, month=month, sort_by="category"):
            if int(row.get("account_id") or 0) not in internal_account_ids:
                continue
            category_id = row.get("category_id")
            cents = int(row.get("expected_amount_cents") or 0)
            if cents == 0:
                continue
            if category_id is not None and int(category_id) in income_category_ids:
                category_name = income_category_names.get(
                    int(category_id),
                    self._dashboard_category_label(row.get("category_name")),
                )
            else:
                category_name = uncategorized_label
            planned_income_by_category[category_name] = planned_income_by_category.get(category_name, 0) + cents

        actual_expense_by_category: dict[str, int] = {}
        actual_income_by_category: dict[str, int] = {}
        relabeled_uncategorized_txn_count = 0
        relabeled_uncategorized_txn_details: list[dict[str, str | int]] = []
        for row in self.context.transactions_service.list_for_month(year=year, month=month, limit=10000):
            if bool(row.get("account_is_external")):
                continue
            txn_type = str(row.get("txn_type") or "").strip().lower()
            if txn_type == "transfer":
                continue
            amount_cents = int(row.get("amount_cents") or 0)
            if amount_cents == 0:
                continue
            category_id = row.get("category_id")
            typed_category_id = int(category_id) if category_id is not None else None
            original_category_name = self._dashboard_category_label(row.get("category_name"))
            txn_id = int(row.get("txn_id") or 0)
            txn_date = str(row.get("txn_date") or "")
            txn_description = self._normalized_display_text(row.get("description") or "")
            if txn_type == "expense" or (txn_type != "income" and amount_cents < 0):
                if typed_category_id is not None and typed_category_id in expense_category_ids:
                    category_name = expense_category_names.get(
                        typed_category_id,
                        self._dashboard_category_label(row.get("category_name")),
                    )
                else:
                    category_name = uncategorized_label
                    relabeled_uncategorized_txn_count += 1
                    if typed_category_id is None:
                        mismatch_reason = "missing category"
                    elif typed_category_id in income_category_ids:
                        mismatch_reason = "category marked as income"
                    else:
                        mismatch_reason = "category not in active expense categories"
                    relabeled_uncategorized_txn_details.append(
                        {
                            "txn_id": txn_id,
                            "txn_date": txn_date,
                            "txn_description": txn_description,
                            "category_id": typed_category_id if typed_category_id is not None else "NULL",
                            "category_name": original_category_name,
                            "txn_type": "expense",
                            "expected_category_type": "expense",
                            "mismatch_reason": mismatch_reason,
                        }
                    )
                actual_expense_by_category[category_name] = (
                    actual_expense_by_category.get(category_name, 0) + abs(amount_cents)
                )
            else:
                if typed_category_id is not None and typed_category_id in income_category_ids:
                    category_name = income_category_names.get(
                        typed_category_id,
                        self._dashboard_category_label(row.get("category_name")),
                    )
                else:
                    category_name = uncategorized_label
                    relabeled_uncategorized_txn_count += 1
                    if typed_category_id is None:
                        mismatch_reason = "missing category"
                    elif typed_category_id in expense_category_ids:
                        mismatch_reason = "category marked as expense"
                    else:
                        mismatch_reason = "category not in active income categories"
                    relabeled_uncategorized_txn_details.append(
                        {
                            "txn_id": txn_id,
                            "txn_date": txn_date,
                            "txn_description": txn_description,
                            "category_id": typed_category_id if typed_category_id is not None else "NULL",
                            "category_name": original_category_name,
                            "txn_type": "income",
                            "expected_category_type": "income",
                            "mismatch_reason": mismatch_reason,
                        }
                    )
                actual_income_by_category[category_name] = (
                    actual_income_by_category.get(category_name, 0) + abs(amount_cents)
                )

        planned_expenses_total = sum(planned_expense_by_category.values())
        actual_expenses_total = sum(actual_expense_by_category.values())
        planned_income_total = sum(planned_income_by_category.values())
        actual_income_total = sum(actual_income_by_category.values())
        end_balance_cents = starting_balance_cents + (actual_income_total - actual_expenses_total)

        expense_rows = self._build_dashboard_category_rows(
            planned_expense_by_category,
            actual_expense_by_category,
            is_income=False,
        )
        income_rows = self._build_dashboard_category_rows(
            planned_income_by_category,
            actual_income_by_category,
            is_income=True,
        )

        return {
            "starting_balance_cents": starting_balance_cents,
            "end_balance_cents": end_balance_cents,
            "planned_expenses_total": planned_expenses_total,
            "actual_expenses_total": actual_expenses_total,
            "planned_income_total": planned_income_total,
            "actual_income_total": actual_income_total,
            "planned_expense_by_category": planned_expense_by_category,
            "actual_expense_by_category": actual_expense_by_category,
            "planned_income_by_category": planned_income_by_category,
            "actual_income_by_category": actual_income_by_category,
            "expense_rows": expense_rows,
            "income_rows": income_rows,
            "relabeled_uncategorized_txn_count": relabeled_uncategorized_txn_count,
            "relabeled_uncategorized_txn_details": relabeled_uncategorized_txn_details,
        }

    def refresh_dashboard(self) -> None:
        month_value = self.dashboard_tab.month_picker.currentText().strip()
        if not month_value:
            return
        try:
            year_str, month_str = month_value.split("-")
            year = int(year_str)
            month = int(month_str)
        except ValueError:
            self.logger.warning("Invalid dashboard month value: %s", month_value)
            return

        snapshot = self._compute_dashboard_snapshot_for_month(year, month)
        starting_balance_cents = int(snapshot["starting_balance_cents"])
        end_balance_cents = int(snapshot["end_balance_cents"])
        planned_expenses_total = int(snapshot["planned_expenses_total"])
        actual_expenses_total = int(snapshot["actual_expenses_total"])
        planned_income_total = int(snapshot["planned_income_total"])
        actual_income_total = int(snapshot["actual_income_total"])
        relabeled_uncategorized_txn_count = int(snapshot["relabeled_uncategorized_txn_count"])
        relabeled_uncategorized_txn_details = list(snapshot.get("relabeled_uncategorized_txn_details") or [])

        self.dashboard_tab.expenses_model.replace_rows(list(snapshot["expense_rows"]))
        self.dashboard_tab.income_model.replace_rows(list(snapshot["income_rows"]))

        self._loading_dashboard_starting_balance = True
        self.dashboard_tab.starting_balance_input.setText(self._format_currency(starting_balance_cents))
        self._loading_dashboard_starting_balance = False
        self.dashboard_tab.end_balance_value.setText(self._format_currency(end_balance_cents))
        self.dashboard_tab.planned_expenses_value.setText(self._format_currency(planned_expenses_total))
        self.dashboard_tab.actual_expenses_value.setText(self._format_currency(actual_expenses_total))
        self.dashboard_tab.planned_income_value.setText(self._format_currency(planned_income_total))
        self.dashboard_tab.actual_income_value.setText(self._format_currency(actual_income_total))

        account_status_rows, total_internal_ending_balance_cents = self._dashboard_account_status_rows(
            year,
            month,
        )
        self.dashboard_tab.account_status_model.replace_rows(account_status_rows)
        self.dashboard_tab.account_status_total_value.setText(
            self._format_currency_balance(total_internal_ending_balance_cents)
        )

        self.logger.info(
            "Dashboard refreshed for %s-%02d (start=%s, end=%s)",
            year,
            month,
            starting_balance_cents,
            end_balance_cents,
        )
        if relabeled_uncategorized_txn_count:
            self.logger.warning(
                "Dashboard relabeled %s transaction(s) to Uncategorized due to category type mismatch for %s-%02d.",
                relabeled_uncategorized_txn_count,
                year,
                month,
            )
            for detail in relabeled_uncategorized_txn_details:
                self.logger.warning(
                    "Dashboard mismatch detail: txn_id=%s, date=%s, txn_type=%s, "
                    "description='%s', category_id=%s, category_name='%s', "
                    "expected_category_type=%s, reason=%s",
                    detail.get("txn_id"),
                    detail.get("txn_date"),
                    detail.get("txn_type"),
                    detail.get("txn_description"),
                    detail.get("category_id"),
                    detail.get("category_name"),
                    detail.get("expected_category_type"),
                    detail.get("mismatch_reason"),
                )

    def save_dashboard_starting_balance(self) -> None:
        if self._loading_dashboard_starting_balance:
            return
        month_value = self.dashboard_tab.month_picker.currentText().strip()
        if not month_value:
            return
        try:
            year_str, month_str = month_value.split("-")
            year = int(year_str)
            month = int(month_str)
        except ValueError:
            self.logger.warning("Invalid dashboard month for starting balance save: %s", month_value)
            return

        try:
            starting_balance_cents = self._parse_currency_cents_allow_negative(
                self.dashboard_tab.starting_balance_input.text()
            )
        except ValueError as exc:
            QMessageBox.warning(self, "Starting Balance", str(exc))
            self.refresh_dashboard()
            return

        self.context.budgeting_service.set_starting_balance(year, month, starting_balance_cents)
        self.statusBar().showMessage(
            f"Saved starting balance for {year:04d}-{month:02d}.",
            3000,
        )
        self.refresh_dashboard()

    def _refresh_transaction_form_choices(self) -> None:
        selected_category_id = self.transactions_tab.category_input.currentData()
        selected_category_text = self._normalized_display_text(self.transactions_tab.category_input.currentText())
        selected_account_id = self._selected_account_id()
        selected_category_type = "income" if self.transactions_tab.income_radio.isChecked() else "expense"

        self.transactions_tab.category_input.clear()
        self.transactions_tab.category_input.addItem("", None)
        for row in self.context.categories_repo.list_active(category_type=selected_category_type):
            self.transactions_tab.category_input.addItem(str(row["name"]), int(row["category_id"]))

        self.account_alias_to_id = {}
        self.account_id_to_alias = {}
        self.account_id_to_type = {}
        self.account_ids_by_type = {}
        account_rows = self.context.accounts_repo.list_active()
        alias_counts: dict[str, int] = {}
        for row in account_rows:
            alias = str(row.get("name") or "").strip()
            if not alias:
                continue
            alias_counts[alias.casefold()] = alias_counts.get(alias.casefold(), 0) + 1

        self.transactions_tab.account_input.clear()
        self.transactions_tab.account_input.addItem("", None)
        for row in account_rows:
            account_alias = str(row.get("name") or "").strip()
            account_id = int(row["account_id"])
            account_type = str(row.get("account_type") or "").strip().lower()
            institution = str(row.get("institution_name") or "").strip()
            alias_key = account_alias.casefold()
            if alias_counts.get(alias_key, 0) > 1 and institution:
                display_label = f"{account_alias} ({institution})"
            else:
                display_label = account_alias

            if account_alias and alias_key not in self.account_alias_to_id:
                self.account_alias_to_id[alias_key] = account_id
            self.account_id_to_alias[account_id] = account_alias
            self.account_id_to_type[account_id] = account_type
            self.account_ids_by_type.setdefault(account_type, []).append(account_id)
            self.transactions_tab.account_input.addItem(display_label, account_id)

        if selected_category_id is not None:
            self._combo_select_data(self.transactions_tab.category_input, selected_category_id)
        elif selected_category_text:
            self.transactions_tab.category_input.setEditText(selected_category_text)
        elif self.transactions_tab.category_input.count() > 0:
            self.transactions_tab.category_input.setCurrentIndex(0)
            self.transactions_tab.category_input.setEditText("")

        if selected_account_id:
            self._combo_select_data(self.transactions_tab.account_input, int(selected_account_id))
        elif self.transactions_tab.account_input.count() > 1:
            self._set_default_account_for_txn_type(selected_category_type)
        else:
            self.transactions_tab.account_input.setCurrentIndex(0)
        self._refresh_transfer_form_choices()

    @staticmethod
    def _combo_select_data(combo, target_data: int | None) -> None:
        if target_data is None:
            if combo.count() > 0:
                combo.setCurrentIndex(0)
                if hasattr(combo, "setEditText"):
                    combo.setEditText("")
            return
        index = combo.findData(target_data)
        if index >= 0:
            combo.setCurrentIndex(index)

    def _clear_transaction_form(self) -> None:
        self.transactions_tab.editing_txn_id = None
        self.transactions_tab.save_button.setEnabled(True)
        self.transactions_tab.delete_button.setEnabled(True)
        self.transactions_tab.txn_date_input.setText(date.today().isoformat())
        self._suppress_type_defaults = True
        self.transactions_tab.expense_radio.setChecked(True)
        self._suppress_type_defaults = False
        self.transactions_tab.amount_input.clear()
        self.transactions_tab.description_input.clear()
        self.transactions_tab.note_input.clear()
        self.transactions_tab.payment_type_input.clear()
        self.transactions_tab.category_input.setCurrentIndex(0)
        self.transactions_tab.category_input.setEditText("")
        self.transactions_tab.subscription_checkbox.setChecked(False)
        self.transactions_tab.subscription_checkbox.setEnabled(True)
        self.transactions_tab.expenses_table.clearSelection()
        self.transactions_tab.income_table.clearSelection()
        self._apply_type_defaults("expense")

    def _refresh_budget_form_choices(self) -> None:
        selected_category_id = self.budget_tab.category_input.currentData()

        self.budget_tab.category_input.clear()
        for row in self.context.categories_repo.list_active(category_type="expense"):
            self.budget_tab.category_input.addItem(str(row["name"]), int(row["category_id"]))

        if selected_category_id is not None:
            self._combo_select_data(self.budget_tab.category_input, int(selected_category_id))
        elif self.budget_tab.category_input.count() > 0:
            self.budget_tab.category_input.setCurrentIndex(0)

    def new_budget_form(self) -> None:
        self.budget_tab.editing_budget_line_id = None
        if self.budget_tab.category_input.count() > 0:
            self.budget_tab.category_input.setCurrentIndex(0)
        self.budget_tab.amount_input.clear()
        self.budget_tab.note_input.clear()
        self.budget_tab.table.clearSelection()

    def _selected_budget_row(self) -> dict | None:
        selection = self.budget_tab.table.selectionModel().selectedRows()
        if not selection:
            return None
        return self.budget_tab.model.row_dict(selection[0].row())

    def on_budget_selection_changed(self) -> None:
        row = self._selected_budget_row()
        if row:
            self._load_budget_into_form(row)

    def _load_budget_into_form(self, row: dict) -> None:
        self.budget_tab.editing_budget_line_id = int(row.get("budget_line_id") or 0)
        self._combo_select_data(self.budget_tab.category_input, row.get("category_id"))
        planned_cents = int(row.get("planned_cents") or 0)
        self.budget_tab.amount_input.setText(f"{planned_cents / 100:.2f}")
        self.budget_tab.note_input.setText(self._normalized_display_text(row.get("note")))

    @staticmethod
    def _build_planned_bills_by_category(rows: list[dict]) -> dict[str, int]:
        planned_by_category: dict[str, int] = {}
        for row in rows:
            category_name = str(row.get("category_name") or "Uncategorized").strip() or "Uncategorized"
            cents = int(row.get("expected_amount_cents") or 0)
            if cents == 0:
                continue
            planned_by_category[category_name] = planned_by_category.get(category_name, 0) + cents
        return planned_by_category

    def refresh_budget_allocations(self) -> None:
        allocation_rows = self.context.budget_allocations_service.list_month_allocations(
            year=self.budget_view_year,
            month=self.budget_view_month,
        )
        planned_bills_by_category = self._build_planned_bills_by_category(
            self.context.bills_service.list_month_bills(
                year=self.budget_view_year,
                month=self.budget_view_month,
                sort_by="category",
            )
        )

        rows: list[dict] = []
        total_allocation_cents = 0
        total_planned_bills_cents = 0
        total_diff_cents = 0
        for allocation in allocation_rows:
            category_name = str(allocation.get("category_name") or "Uncategorized").strip() or "Uncategorized"
            allocation_cents = int(allocation.get("planned_cents") or 0)
            planned_bills_cents = int(planned_bills_by_category.get(category_name, 0))
            diff_cents = allocation_cents - planned_bills_cents
            total_allocation_cents += allocation_cents
            total_planned_bills_cents += planned_bills_cents
            total_diff_cents += diff_cents
            rows.append(
                {
                    **allocation,
                    "allocation_display": self._format_currency(allocation_cents),
                    "planned_bills_display": self._format_currency(planned_bills_cents),
                    "diff_display": self._format_currency_signed(diff_cents),
                    "note": str(allocation.get("note") or ""),
                }
            )

        rows.sort(
            key=lambda r: (
                str(r.get("category_name") or "").casefold(),
                int(r.get("budget_line_id") or 0),
            )
        )
        planned_income_total_cents = sum(
            int(row.get("expected_amount_cents") or 0)
            for row in self.context.income_service.list_month_income(
                year=self.budget_view_year,
                month=self.budget_view_month,
                sort_by="category",
            )
        )
        unallocated_cents = planned_income_total_cents - total_allocation_cents
        self.budget_tab.model.replace_rows(rows)
        self.budget_tab.totals_model.replace_rows(
            [
                {
                    "metric": "Allocation",
                    "value_display": self._format_currency(total_allocation_cents),
                },
                {
                    "metric": "Planned",
                    "value_display": self._format_currency(total_planned_bills_cents),
                },
                {
                    "metric": "Diff",
                    "value_display": self._format_currency_signed(total_diff_cents),
                },
                {
                    "metric": "Total Planned Income",
                    "value_display": self._format_currency(planned_income_total_cents),
                },
                {
                    "metric": "Unallocated",
                    "value_display": self._format_currency_signed(unallocated_cents),
                },
            ]
        )
        self.logger.info(
            "Loaded %s budget allocations for %s-%02d",
            len(rows),
            self.budget_view_year,
            self.budget_view_month,
        )
        self.refresh_dashboard()

    def save_budget_allocation(self) -> None:
        category_id = self.budget_tab.category_input.currentData()
        if category_id is None:
            QMessageBox.warning(self, "Save Budget Allocation", "Category is required.")
            return
        try:
            parsed = self._parse_currency_cents_or_none(self.budget_tab.amount_input.text())
            planned_cents = int(parsed or 0)
            note = self._normalized_display_text(self.budget_tab.note_input.text()) or None

            if self.budget_tab.editing_budget_line_id:
                updated = self.context.budget_allocations_service.update_month_allocation(
                    budget_line_id=int(self.budget_tab.editing_budget_line_id),
                    category_id=int(category_id),
                    planned_cents=planned_cents,
                    note=note,
                )
                if not updated:
                    QMessageBox.warning(
                        self,
                        "Save Budget Allocation",
                        "Selected budget allocation no longer exists.",
                    )
                    self.refresh_budget_allocations()
                    self.new_budget_form()
                    return
            else:
                self.context.budget_allocations_service.upsert_month_allocation(
                    year=self.budget_view_year,
                    month=self.budget_view_month,
                    category_id=int(category_id),
                    planned_cents=planned_cents,
                    note=note,
                )
            self._mark_budget_month_dirty(self.budget_view_year, self.budget_view_month)
        except ValueError as exc:
            QMessageBox.warning(self, "Save Budget Allocation", str(exc))
            return
        except Exception as exc:
            self.logger.error("Save budget allocation failed: %s", exc)
            QMessageBox.critical(self, "Save Budget Allocation Failed", str(exc))
            return

        self.refresh_budget_allocations()
        self.new_budget_form()

    def delete_budget_allocation(self) -> None:
        row = self._selected_budget_row()
        if not row:
            QMessageBox.information(self, "Delete Budget Allocation", "Select an allocation row to delete.")
            return
        budget_line_id = int(row.get("budget_line_id") or 0)
        if not budget_line_id:
            QMessageBox.warning(self, "Delete Budget Allocation", "Invalid allocation selection.")
            return
        answer = QMessageBox.question(
            self,
            "Delete Budget Allocation",
            f"Delete allocation for '{row.get('category_name', '')}'?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if answer != QMessageBox.Yes:
            return

        try:
            deleted = self.context.budget_allocations_service.delete_month_allocation(budget_line_id)
        except Exception as exc:
            self.logger.error("Delete budget allocation failed: %s", exc)
            QMessageBox.critical(self, "Delete Budget Allocation Failed", str(exc))
            return
        if not deleted:
            QMessageBox.warning(self, "Delete Budget Allocation", "Selected allocation no longer exists.")
            return

        self._mark_budget_month_dirty(self.budget_view_year, self.budget_view_month)
        self.refresh_budget_allocations()
        self.new_budget_form()

    def set_budget_sort(self, sort_key: str) -> None:
        self.budget_sort_key = sort_key
        self.refresh_budget_allocations()

    def refresh_budget_for_selected_month(self) -> None:
        month_label = f"{self.budget_view_year:04d}-{self.budget_view_month:02d}"
        if self._is_budget_month_dirty(self.budget_view_year, self.budget_view_month):
            answer = QMessageBox.warning(
                self,
                "Refresh Category Budgets",
                "Refreshing category budgets will replace all monthly allocation instances for "
                f"{month_label} with current global budget category definitions.\n\n"
                "Any edits made to monthly budget allocations for this month will be lost.\n\n"
                "Do you want to continue?",
                QMessageBox.Yes | QMessageBox.Cancel,
                QMessageBox.Cancel,
            )
            if answer != QMessageBox.Yes:
                self.statusBar().showMessage("Refresh Category Budgets canceled.", 3000)
                return

        deleted, inserted = self.context.budget_allocations_service.regenerate_for_month(
            self.budget_view_year,
            self.budget_view_month,
        )
        self._clear_budget_month_dirty(self.budget_view_year, self.budget_view_month)
        self.refresh_budget_allocations()
        self.statusBar().showMessage(
            f"Refreshed category budgets for {month_label}: replaced {deleted}, generated {inserted}.",
            5000,
        )

    def _refresh_bill_form_choices(self) -> None:
        selected_category_id = self.bills_tab.category_input.currentData()

        self.bills_tab.category_input.clear()
        self.bills_tab.category_input.addItem("", None)
        for row in self.context.categories_repo.list_active(category_type="expense"):
            self.bills_tab.category_input.addItem(str(row["name"]), int(row["category_id"]))

        if selected_category_id is not None:
            self._combo_select_data(self.bills_tab.category_input, selected_category_id)
        else:
            self.bills_tab.category_input.setCurrentIndex(0)

    def new_bill_form(self) -> None:
        self.bills_tab.editing_bill_id = None
        self._editing_bill_source_system = None
        self.bills_tab.bill_name_input.clear()
        self.bills_tab.start_date_input.setText(date.today().isoformat())
        self.bills_tab.date_paid_input.clear()
        self.bills_tab.interval_count_input.clear()
        self.bills_tab.interval_unit_combo.setCurrentText("months")
        self.bills_tab.amount_input.clear()
        self.bills_tab.note_input.clear()
        self.bills_tab.category_input.setCurrentIndex(0)
        self.bills_tab.table.clearSelection()

    def _selected_bill_row(self) -> dict | None:
        selection = self.bills_tab.table.selectionModel().selectedRows()
        if not selection:
            return None
        return self.bills_tab.model.row_dict(selection[0].row())

    def on_bill_selection_changed(self) -> None:
        row = self._selected_bill_row()
        if row:
            self._load_bill_into_form(row)

    def _load_bill_into_form(self, row: dict) -> None:
        self.bills_tab.editing_bill_id = int(row.get("bill_occurrence_id") or 0)
        self._editing_bill_source_system = self._normalized_source_system(
            str(row.get("source_system") or "")
        )
        self.bills_tab.bill_name_input.setText(self._normalized_display_text(row.get("name")))
        due_date = self._normalized_display_text(row.get("payment_due") or row.get("expected_date"))
        self.bills_tab.start_date_input.setText(due_date or date.today().isoformat())
        self.bills_tab.date_paid_input.setText(self._normalized_display_text(row.get("paid_date")))
        self.bills_tab.interval_count_input.setText(str(int(row.get("interval_count") or 1)))
        raw_unit = self._normalized_display_text(row.get("interval_unit")) or "months"
        if raw_unit.endswith("s"):
            normalized_unit = raw_unit
        else:
            normalized_unit = f"{raw_unit}s"
        if self.bills_tab.interval_unit_combo.findText(normalized_unit) >= 0:
            self.bills_tab.interval_unit_combo.setCurrentText(normalized_unit)
        else:
            self.bills_tab.interval_unit_combo.setCurrentText("months")
        amount_cents = row.get("expected_amount_cents")
        if amount_cents is None:
            self.bills_tab.amount_input.clear()
        else:
            self.bills_tab.amount_input.setText(f"{int(amount_cents) / 100:.2f}")
        self.bills_tab.note_input.setText(self._normalized_display_text(row.get("notes")))
        self._combo_select_data(self.bills_tab.category_input, row.get("category_id"))

    @staticmethod
    def _parse_currency_cents_or_none(amount_text: str) -> int | None:
        cleaned = amount_text.strip().replace("$", "").replace(",", "")
        if not cleaned:
            return None
        try:
            amount = Decimal(cleaned)
        except InvalidOperation as exc:
            raise ValueError("Amount must be numeric (example: 74.68).") from exc
        cents = int((amount * 100).quantize(Decimal("1"), rounding=ROUND_HALF_UP))
        if cents < 0:
            raise ValueError("Amount cannot be negative.")
        return cents

    def _build_bill_occurrence_payload(self) -> dict:
        due_date_text = self.bills_tab.start_date_input.text().strip()
        try:
            datetime.strptime(due_date_text, "%Y-%m-%d")
        except ValueError as exc:
            raise ValueError("Payment Due must be in YYYY-MM-DD format.") from exc

        paid_date_text = self.bills_tab.date_paid_input.text().strip()
        if paid_date_text:
            try:
                datetime.strptime(paid_date_text, "%Y-%m-%d")
            except ValueError as exc:
                raise ValueError("Date Paid must be in YYYY-MM-DD format.") from exc

        amount_cents = self._parse_currency_cents_or_none(self.bills_tab.amount_input.text())

        return {
            "expected_date": due_date_text,
            "paid_date": paid_date_text or None,
            "expected_amount_cents": amount_cents,
            "notes": self._normalized_display_text(self.bills_tab.note_input.text()) or None,
        }

    def save_bill(self) -> None:
        bill_id = self.bills_tab.editing_bill_id
        action_label = "Save Bill Update"
        if not bill_id:
            QMessageBox.information(
                self,
                action_label,
                "Select a bill row to update this month occurrence.",
            )
            return

        try:
            payload = self._build_bill_occurrence_payload()
            updated = self.context.bills_service.update_occurrence(
                bill_occurrence_id=bill_id,
                expected_date=payload["expected_date"],
                expected_amount_cents=payload["expected_amount_cents"],
                paid_date=payload["paid_date"],
                note=payload["notes"],
            )
            if not updated:
                QMessageBox.warning(self, action_label, "Selected bill row no longer exists.")
                self.refresh_bills()
                self.new_bill_form()
                return
            self._mark_bills_month_dirty(
                self.bills_view_year,
                self.bills_view_month,
                self._editing_bill_source_system,
            )
            self.logger.info("Updated bill occurrence %s", bill_id)
            self.statusBar().showMessage("Bill updated for selected month.", 3000)
        except ValueError as exc:
            QMessageBox.warning(self, action_label, str(exc))
            return
        except Exception as exc:
            self.logger.error("Save bill failed: %s", exc)
            QMessageBox.critical(self, "Save Bill Failed", str(exc))
            return

        self.refresh_bills()
        self.new_bill_form()

    def delete_bill(self) -> None:
        row = self._selected_bill_row()
        if not row:
            QMessageBox.information(self, "Delete Bill", "Select a bill to delete.")
            return
        bill_id = int(row.get("bill_occurrence_id") or 0)
        if not bill_id:
            QMessageBox.warning(self, "Delete Bill", "Invalid bill selection.")
            return

        answer = QMessageBox.question(
            self,
            "Delete Bill",
            f"Delete bill '{row.get('name', '')}' for this month?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if answer != QMessageBox.Yes:
            return

        try:
            deleted = self.context.bills_service.delete_occurrence(bill_id)
        except Exception as exc:
            self.logger.error("Delete bill failed: %s", exc)
            QMessageBox.critical(self, "Delete Bill Failed", str(exc))
            return

        if not deleted:
            QMessageBox.warning(self, "Delete Bill", "Selected bill no longer exists.")
        else:
            self._mark_bills_month_dirty(
                self.bills_view_year,
                self.bills_view_month,
                str(row.get("source_system") or ""),
            )
            self.logger.info("Deleted bill occurrence %s", bill_id)
            self.statusBar().showMessage("Bill deleted from selected month.", 3000)

        self.refresh_bills()
        self.new_bill_form()

    def set_bills_sort(self, sort_key: str) -> None:
        self.bills_sort_key = sort_key
        self.refresh_bills()

    def _refresh_income_form_choices(self) -> None:
        selected_category_id = self.income_tab.category_input.currentData()
        selected_account_name = self._normalized_display_text(self.income_tab.account_input.text())

        self.income_tab.category_input.clear()
        self.income_tab.category_input.addItem("", None)
        for row in self.context.categories_repo.list_active(category_type="income"):
            self.income_tab.category_input.addItem(str(row["name"]), int(row["category_id"]))
        if selected_category_id is not None:
            self._combo_select_data(self.income_tab.category_input, selected_category_id)
        else:
            self.income_tab.category_input.setCurrentIndex(0)

        account_name = selected_account_name
        if not account_name:
            accounts = self.context.accounts_repo.list_active()
            if accounts:
                account_name = str(accounts[0]["name"])
        self.income_tab.account_input.setText(account_name)

    def new_income_form(self) -> None:
        self.income_tab.editing_income_occurrence_id = None
        self.income_tab.description_input.clear()
        self.income_tab.start_date_input.setText(date.today().isoformat())
        self.income_tab.interval_count_input.clear()
        self.income_tab.interval_unit_combo.setCurrentText("months")
        self.income_tab.amount_input.clear()
        self.income_tab.note_input.clear()
        self.income_tab.category_input.setCurrentIndex(0)
        self.income_tab.account_input.clear()
        self.income_tab.table.clearSelection()

    def _selected_income_row(self) -> dict | None:
        selection = self.income_tab.table.selectionModel().selectedRows()
        if not selection:
            return None
        return self.income_tab.model.row_dict(selection[0].row())

    def on_income_selection_changed(self) -> None:
        row = self._selected_income_row()
        if row:
            self._load_income_into_form(row)

    def _load_income_into_form(self, row: dict) -> None:
        self.income_tab.editing_income_occurrence_id = int(row.get("income_occurrence_id") or 0)
        self.income_tab.description_input.setText(self._normalized_display_text(row.get("description")))
        due_date = self._normalized_display_text(row.get("payment_due") or row.get("expected_date"))
        self.income_tab.start_date_input.setText(due_date or date.today().isoformat())
        self.income_tab.interval_count_input.setText(str(int(row.get("interval_count") or 1)))
        raw_unit = self._normalized_display_text(row.get("interval_unit")) or "months"
        normalized_unit = raw_unit if raw_unit.endswith("s") or raw_unit == "once" else f"{raw_unit}s"
        if self.income_tab.interval_unit_combo.findText(normalized_unit) >= 0:
            self.income_tab.interval_unit_combo.setCurrentText(normalized_unit)
        else:
            self.income_tab.interval_unit_combo.setCurrentText("months")
        amount_cents = row.get("expected_amount_cents")
        if amount_cents is None:
            self.income_tab.amount_input.clear()
        else:
            self.income_tab.amount_input.setText(f"{int(amount_cents) / 100:.2f}")
        self.income_tab.note_input.setText(self._normalized_display_text(row.get("notes")))
        self._combo_select_data(self.income_tab.category_input, row.get("category_id"))
        self.income_tab.account_input.setText(self._normalized_display_text(row.get("account_name")))

    def _build_income_occurrence_payload(self) -> dict:
        due_date_text = self.income_tab.start_date_input.text().strip()
        try:
            datetime.strptime(due_date_text, "%Y-%m-%d")
        except ValueError as exc:
            raise ValueError("Payment Due must be in YYYY-MM-DD format.") from exc
        amount_cents = self._parse_currency_cents_or_none(self.income_tab.amount_input.text())
        return {
            "expected_date": due_date_text,
            "expected_amount_cents": amount_cents,
            "notes": self._normalized_display_text(self.income_tab.note_input.text()) or None,
        }

    def save_income(self) -> None:
        income_occurrence_id = self.income_tab.editing_income_occurrence_id
        if not income_occurrence_id:
            QMessageBox.information(
                self,
                "Save Income Update",
                "Select an income row to update this month occurrence.",
            )
            return
        try:
            payload = self._build_income_occurrence_payload()
            updated = self.context.income_service.update_occurrence(
                income_occurrence_id=income_occurrence_id,
                expected_date=payload["expected_date"],
                expected_amount_cents=payload["expected_amount_cents"],
                note=payload["notes"],
            )
            if not updated:
                QMessageBox.warning(self, "Save Income Update", "Selected income row no longer exists.")
                self.refresh_income()
                self.new_income_form()
                return
            self._mark_income_month_dirty(self.income_view_year, self.income_view_month)
            self.logger.info("Updated income occurrence %s", income_occurrence_id)
            self.statusBar().showMessage("Income updated for selected month.", 3000)
        except ValueError as exc:
            QMessageBox.warning(self, "Save Income Update", str(exc))
            return
        except Exception as exc:
            self.logger.error("Save income failed: %s", exc)
            QMessageBox.critical(self, "Save Income Failed", str(exc))
            return
        self.refresh_income()
        self.new_income_form()

    def delete_income(self) -> None:
        row = self._selected_income_row()
        if not row:
            QMessageBox.information(self, "Delete Income", "Select an income row to delete.")
            return
        income_occurrence_id = int(row.get("income_occurrence_id") or 0)
        if not income_occurrence_id:
            QMessageBox.warning(self, "Delete Income", "Invalid income selection.")
            return
        answer = QMessageBox.question(
            self,
            "Delete Income",
            f"Delete income '{row.get('description', '')}' for this month?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if answer != QMessageBox.Yes:
            return
        try:
            deleted = self.context.income_service.delete_occurrence(income_occurrence_id)
        except Exception as exc:
            self.logger.error("Delete income failed: %s", exc)
            QMessageBox.critical(self, "Delete Income Failed", str(exc))
            return
        if not deleted:
            QMessageBox.warning(self, "Delete Income", "Selected income row no longer exists.")
        else:
            self._mark_income_month_dirty(self.income_view_year, self.income_view_month)
            self.logger.info("Deleted income occurrence %s", income_occurrence_id)
            self.statusBar().showMessage("Income deleted from selected month.", 3000)
        self.refresh_income()
        self.new_income_form()

    def set_income_sort(self, sort_key: str) -> None:
        self.income_sort_key = sort_key
        self.refresh_income()

    @staticmethod
    def _build_income_category_subtotals(rows: list[dict]) -> list[dict]:
        totals_by_category: dict[str, int] = {}
        for row in rows:
            category_name = str(row.get("category_name") or "Uncategorized").strip() or "Uncategorized"
            amount_cents = int(row.get("expected_amount_cents") or 0)
            totals_by_category[category_name] = totals_by_category.get(category_name, 0) + amount_cents
        subtotal_rows = [
            {
                "category_name": category_name,
                "subtotal_display": f"${cents / 100:.2f}",
            }
            for category_name, cents in totals_by_category.items()
        ]
        subtotal_rows.sort(
            key=lambda item: (
                str(item.get("category_name") or "").strip().casefold(),
                str(item.get("category_name") or "").strip(),
            )
        )
        return subtotal_rows

    def refresh_income(self) -> None:
        rows = self.context.income_service.list_month_income(
            year=self.income_view_year,
            month=self.income_view_month,
            sort_by=self.income_sort_key,
        )
        income_category_rows = self.context.categories_repo.list_active(category_type="income")
        income_category_ids = {int(row["category_id"]) for row in income_category_rows}
        income_category_names = {
            int(row["category_id"]): self._dashboard_category_label(row.get("name"))
            for row in income_category_rows
        }
        normalized_rows: list[dict] = []
        relabeled_uncategorized_count = 0
        for row in rows:
            data = dict(row)
            category_id = data.get("category_id")
            if category_id is not None and int(category_id) in income_category_ids:
                data["category_name"] = income_category_names.get(
                    int(category_id),
                    self._dashboard_category_label(data.get("category_name")),
                )
            else:
                data["category_name"] = "Uncategorized"
                relabeled_uncategorized_count += 1
            normalized_rows.append(data)

        self.income_tab.model.replace_rows(normalized_rows)
        subtotal_rows = self._build_income_category_subtotals(normalized_rows)
        self.income_tab.category_totals_model.replace_rows(subtotal_rows)
        total_cents = sum(int(row.get("expected_amount_cents") or 0) for row in rows)
        self.income_tab.total_income_value_label.setText(f"${total_cents / 100:.2f}")
        self.logger.info(
            "Loaded %s income occurrences for %s-%02d (sort=%s)",
            len(rows),
            self.income_view_year,
            self.income_view_month,
            self.income_sort_key,
        )
        if relabeled_uncategorized_count:
            self.logger.warning(
                "Income view relabeled %s row(s) to Uncategorized due to category type mismatch for %s-%02d.",
                relabeled_uncategorized_count,
                self.income_view_year,
                self.income_view_month,
            )
        self.refresh_dashboard()

    def refresh_income_for_selected_month(self) -> None:
        month_label = f"{self.income_view_year:04d}-{self.income_view_month:02d}"
        if self._is_income_month_dirty(self.income_view_year, self.income_view_month):
            answer = QMessageBox.warning(
                self,
                "Refresh Income",
                "Refreshing income will replace all income instances for "
                f"{month_label} with current global definitions.\n\n"
                "Any edits made to monthly income instances for this month will be lost.\n\n"
                "Do you want to continue?",
                QMessageBox.Yes | QMessageBox.Cancel,
                QMessageBox.Cancel,
            )
            if answer != QMessageBox.Yes:
                self.statusBar().showMessage("Refresh Income canceled.", 3000)
                return

        deleted, inserted = self.context.income_service.regenerate_for_month(
            self.income_view_year,
            self.income_view_month,
        )
        self._clear_income_month_dirty(self.income_view_year, self.income_view_month)
        self.refresh_income()
        self.logger.info(
            "Refreshed income for %s-%02d: replaced %s occurrences, generated %s from definitions.",
            self.income_view_year,
            self.income_view_month,
            deleted,
            inserted,
        )
        self.statusBar().showMessage(
            f"Refreshed income for {month_label}: replaced {deleted}, generated {inserted}.",
            5000,
        )

    def _selected_account_id(self) -> int | None:
        selected = self.transactions_tab.account_input.currentData()
        if selected is None:
            return None
        try:
            value = int(selected)
        except (TypeError, ValueError):
            return None
        return value if value > 0 else None

    def _set_account_by_id(self, account_id: int | None) -> bool:
        if account_id is None:
            return False
        index = self.transactions_tab.account_input.findData(int(account_id))
        if index <= 0:
            return False
        self.transactions_tab.account_input.setCurrentIndex(index)
        return True

    def _set_default_account_for_txn_type(self, txn_type: str) -> None:
        preferred_types = (
            ("checking", "savings", "cash", "credit")
            if txn_type == "income"
            else ("credit", "checking", "cash", "savings")
        )
        for account_type in preferred_types:
            ids = self.account_ids_by_type.get(account_type, [])
            if ids and self._set_account_by_id(ids[0]):
                return
        if self.transactions_tab.account_input.count() > 1:
            self.transactions_tab.account_input.setCurrentIndex(1)

    def _on_type_changed(self, txn_type: str, checked: bool) -> None:
        if not checked or self._suppress_type_defaults:
            return
        self._apply_type_defaults(txn_type)
        self._refresh_transaction_form_choices()

    def _apply_type_defaults(self, txn_type: str) -> None:
        if txn_type == "income":
            self._set_default_account_for_txn_type("income")
            self.transactions_tab.subscription_checkbox.setChecked(False)
            self.transactions_tab.subscription_checkbox.setEnabled(False)
            self.transactions_tab.tax_checkbox.setChecked(True)
            return

        self._set_default_account_for_txn_type("expense")
        self.transactions_tab.subscription_checkbox.setEnabled(True)
        self.transactions_tab.subscription_checkbox.setChecked(False)
        self.transactions_tab.tax_checkbox.setChecked(False)

    def _selected_transaction_row(self) -> dict | None:
        expense_selection = self.transactions_tab.expenses_table.selectionModel().selectedRows()
        income_selection = self.transactions_tab.income_table.selectionModel().selectedRows()
        if expense_selection:
            row = self.transactions_tab.expense_model.row_dict(expense_selection[0].row())
            if row and int(row.get("txn_id") or 0) > 0:
                return row
            return None
        if income_selection:
            row = self.transactions_tab.income_model.row_dict(income_selection[0].row())
            if row and int(row.get("txn_id") or 0) > 0:
                return row
            return None
        return None

    def _on_expense_row_clicked(self, index) -> None:
        if self._suppress_selection_autoload:
            return
        self.transactions_tab.income_table.clearSelection()
        row = self.transactions_tab.expense_model.row_dict(index.row())
        if row and int(row.get("txn_id") or 0) > 0:
            self._show_transfer_rule_read_only_warning(row)
            self._load_transaction_into_form(row, show_status=False)

    def _on_income_row_clicked(self, index) -> None:
        if self._suppress_selection_autoload:
            return
        self.transactions_tab.expenses_table.clearSelection()
        row = self.transactions_tab.income_model.row_dict(index.row())
        if row and int(row.get("txn_id") or 0) > 0:
            self._show_transfer_rule_read_only_warning(row)
            self._load_transaction_into_form(row, show_status=False)

    def on_transaction_selection_changed(self) -> None:
        if self._suppress_selection_autoload:
            return
        row = self._selected_transaction_row()
        if row:
            self._load_transaction_into_form(row, show_status=False)

    def _show_transfer_rule_read_only_warning(self, row: dict) -> None:
        txn_type = str(row.get("txn_type") or "").strip().lower()
        source_system = str(row.get("source_system") or "").strip().lower()
        if txn_type != "transfer" or source_system != "xlsx_import":
            return
        QMessageBox.information(
            self,
            "Rule-Based Transfer (Read-Only)",
            "This transaction is a rule-based transfer generated during import.\n"
            "Transfer rule definitions can override spreadsheet account values.\n\n"
            "To fix this, edit the transfer rule in Settings.\n"
            "Manual editing on the Transactions tab is disabled for transfer rows.",
        )

    def _load_transaction_into_form(self, row: dict, show_status: bool = True) -> None:
        txn_id = int(row.get("txn_id") or 0)
        if txn_id <= 0:
            return
        self.transactions_tab.editing_txn_id = txn_id
        amount_cents = int(row.get("amount_cents") or 0)
        txn_type = str(row.get("txn_type") or "").strip().lower()
        self.transactions_tab.txn_date_input.setText(str(row.get("txn_date", "")))
        self._suppress_type_defaults = True
        if amount_cents > 0:
            self.transactions_tab.income_radio.setChecked(True)
        else:
            self.transactions_tab.expense_radio.setChecked(True)
        self._suppress_type_defaults = False
        self._refresh_transaction_form_choices()
        self.transactions_tab.amount_input.setText(f"{abs(amount_cents) / 100:.2f}")
        self.transactions_tab.description_input.setText(
            self._normalized_display_text(row.get("description"))
        )
        self.transactions_tab.note_input.setText(
            self._normalized_display_text(row.get("note"))
        )
        self.transactions_tab.payment_type_input.setText(
            self._normalized_display_text(row.get("payment_type"))
        )
        self.transactions_tab.subscription_checkbox.setChecked(bool(row.get("is_subscription")))
        self.transactions_tab.subscription_checkbox.setEnabled(amount_cents <= 0)
        self.transactions_tab.tax_checkbox.setChecked(bool(row.get("tax_deductible")))
        self._combo_select_data(self.transactions_tab.category_input, row.get("category_id"))
        self._combo_select_data(
            self.transactions_tab.account_input,
            row.get("account_id"),
        )
        is_transfer = txn_type == "transfer"
        self.transactions_tab.save_button.setEnabled(not is_transfer)
        self.transactions_tab.delete_button.setEnabled(not is_transfer)
        if is_transfer:
            self.statusBar().showMessage(
                "Transfer actions are read-only here. Use the Transfers tab to manage them.",
                3500,
            )
        if is_transfer and show_status:
            return
        if show_status:
            self.statusBar().showMessage("Loaded selected transaction into editor.", 3000)

    @staticmethod
    def _transactions_spacer_row() -> dict:
        return {
            "_is_spacer_row": True,
            "txn_id": 0,
            "txn_date": "",
            "description_display": "",
            "display_amount_cents": 0,
            "category_name": "",
            "account_name": "",
            "payment_type": "",
            "is_subscription": False,
            "tax_deductible": False,
        }

    @staticmethod
    def _parse_amount_cents(amount_text: str, txn_type_text: str) -> int:
        cleaned = amount_text.strip().replace("$", "").replace(",", "")
        if not cleaned:
            raise ValueError("Amount is required.")
        try:
            amount = Decimal(cleaned)
        except InvalidOperation as exc:
            raise ValueError("Amount must be numeric (example: 74.68).") from exc

        cents = int((amount * 100).quantize(Decimal("1"), rounding=ROUND_HALF_UP))
        if cents <= 0:
            raise ValueError("Amount must be greater than 0.")
        if txn_type_text.lower() == "expense":
            return -abs(cents)
        return abs(cents)

    def _build_transaction_input(self, existing: dict | None = None) -> TransactionInput:
        date_text = self.transactions_tab.txn_date_input.text().strip()
        if not date_text:
            raise ValueError("Date is required.")
        try:
            datetime.strptime(date_text, "%Y-%m-%d")
        except ValueError as exc:
            raise ValueError("Date must be in yyyy-mm-dd format.") from exc

        txn_type_text = "Income" if self.transactions_tab.income_radio.isChecked() else "Expense"
        amount_cents = self._parse_amount_cents(self.transactions_tab.amount_input.text(), txn_type_text)
        description = self._normalized_display_text(self.transactions_tab.description_input.text())
        note_text = self._normalized_display_text(self.transactions_tab.note_input.text())
        payment_type = self._normalized_display_text(self.transactions_tab.payment_type_input.text())

        category_id = self.transactions_tab.category_input.currentData()
        if category_id is None:
            category_name = self.transactions_tab.category_input.currentText().strip()
            if not category_name:
                raise ValueError("Category is required.")
            category_type = "income" if txn_type_text.lower() == "income" else "expense"
            existing_category = self.context.categories_repo.find_by_name(
                category_name,
                category_type=category_type,
            )
            if existing_category:
                category_id = int(existing_category["category_id"])
            else:
                category_id = self.context.categories_repo.upsert(
                    category_name,
                    is_income=(txn_type_text.lower() == "income"),
                )

        account_id = self._selected_account_id()
        if account_id is None:
            raise ValueError("Account is required.")

        internal_payee = description or "Transaction"

        source_system = str((existing or {}).get("source_system") or "manual")
        source_uid = (existing or {}).get("source_uid") or f"manual:{uuid.uuid4()}"
        existing_import_period = str((existing or {}).get("import_period_key") or "").strip() or None
        if source_system == "xlsx_import":
            import_period_key = existing_import_period or date_text[:7]
        else:
            import_period_key = date_text[:7]
        tax_deductible = self.transactions_tab.tax_checkbox.isChecked()
        is_subscription = (
            self.transactions_tab.subscription_checkbox.isChecked()
            if txn_type_text.lower() == "expense"
            else False
        )
        tax_category = (existing or {}).get("tax_category")
        if tax_deductible and txn_type_text.lower() == "expense" and not tax_category:
            tax_category = "Other"
        if not tax_deductible or txn_type_text.lower() == "income":
            tax_category = None

        return TransactionInput(
            txn_date=date_text,
            amount_cents=amount_cents,
            txn_type="income" if amount_cents > 0 else "expense",
            payee=internal_payee,
            description=description or None,
            category_id=int(category_id),
            account_id=int(account_id),
            is_subscription=is_subscription,
            note=note_text or None,
            source_system=source_system,
            source_uid=str(source_uid),
            import_period_key=import_period_key,
            payment_type=payment_type or None,
            import_hash=(existing or {}).get("import_hash"),
            tax_deductible=tax_deductible,
            tax_category=tax_category,
            tax_note=(existing or {}).get("tax_note"),
            receipt_uri=(existing or {}).get("receipt_uri"),
            transfer_group_id=(existing or {}).get("transfer_group_id"),
        )

    @staticmethod
    def _normalized_display_text(value: object | None) -> str:
        text = str(value).strip() if value is not None else ""
        if text in {"...", "…"}:
            return ""
        return text

    def add_transaction(self) -> None:
        self.save_transaction(force_insert=True)

    def _transfer_account_choices(self) -> list[dict]:
        rows = self.context.accounts_repo.list_active()
        choices: list[dict] = []
        for row in rows:
            institution_name = str(row.get("institution_name") or "").strip()
            account_name = str(row.get("name") or "").strip()
            account_type = str(row.get("account_type") or "").strip().lower()
            label = f"{institution_name} • {account_name} ({account_type})" if institution_name else account_name
            choices.append(
                {
                    "account_id": int(row["account_id"]),
                    "account_type": account_type,
                    "display_label": label,
                    "name": account_name,
                    "institution_name": institution_name,
                }
            )
        return choices

    @staticmethod
    def _find_default_account_id(choices: list[dict], account_type: str) -> int | None:
        target_type = str(account_type).strip().lower()
        for row in choices:
            if str(row.get("account_type") or "").strip().lower() == target_type:
                return int(row["account_id"])
        return int(choices[0]["account_id"]) if choices else None

    def _refresh_transfer_form_choices(self) -> None:
        from_combo = self.transfers_tab.transfer_from_account_combo
        to_combo = self.transfers_tab.transfer_to_account_combo
        account_choices = self._transfer_account_choices()
        selected_from_id = int(from_combo.currentData() or 0)
        selected_to_id = int(to_combo.currentData() or 0)

        from_combo.blockSignals(True)
        to_combo.blockSignals(True)
        try:
            from_combo.clear()
            to_combo.clear()
            from_combo.addItem("", None)
            to_combo.addItem("", None)
            for row in account_choices:
                from_combo.addItem(
                    str(row["display_label"]), int(row["account_id"])
                )
                to_combo.addItem(
                    str(row["display_label"]), int(row["account_id"])
                )

            from_idx = from_combo.findData(selected_from_id)
            to_idx = to_combo.findData(selected_to_id)
            if from_idx > 0:
                from_combo.setCurrentIndex(from_idx)
            if to_idx > 0:
                to_combo.setCurrentIndex(to_idx)
        finally:
            from_combo.blockSignals(False)
            to_combo.blockSignals(False)

    def _build_transfer_input_from_form(self) -> TransferInput:
        txn_date = self.transfers_tab.transfer_date_input.text().strip()
        if not txn_date:
            raise ValueError("Transfer date is required.")
        datetime.strptime(txn_date, "%Y-%m-%d")
        amount_cents = self._parse_amount_cents(self.transfers_tab.transfer_amount_input.text(), "Transfer")
        amount_cents = abs(amount_cents)
        if amount_cents <= 0:
            raise ValueError("Transfer amount must be greater than 0.")

        from_account_id = int(self.transfers_tab.transfer_from_account_combo.currentData() or 0)
        to_account_id = int(self.transfers_tab.transfer_to_account_combo.currentData() or 0)
        if from_account_id <= 0 or to_account_id <= 0:
            raise ValueError("Transfer requires both From and To accounts.")
        if from_account_id == to_account_id:
            raise ValueError("From and To accounts must be different.")

        description = self._normalized_display_text(self.transfers_tab.transfer_description_input.text())
        note = self._normalized_display_text(self.transfers_tab.transfer_note_input.text())
        payee = description or "Transfer"
        return TransferInput(
            txn_date=txn_date,
            amount_cents=amount_cents,
            from_account_id=from_account_id,
            to_account_id=to_account_id,
            payee=payee,
            description=description or None,
            note=note or None,
            source_system="manual",
            import_period_key=txn_date[:7],
        )

    def new_transfer_form(self) -> None:
        self.editing_transfer_group_id = None
        self.editing_transfer_source_system = None
        self.transfers_tab.transfers_table.clearSelection()
        for widget in (
            self.transfers_tab.transfer_date_input,
            self.transfers_tab.transfer_amount_input,
            self.transfers_tab.transfer_from_account_combo,
            self.transfers_tab.transfer_to_account_combo,
            self.transfers_tab.transfer_description_input,
            self.transfers_tab.transfer_note_input,
        ):
            widget.setEnabled(True)
        self.transfers_tab.transfer_date_input.clear()
        self.transfers_tab.transfer_amount_input.clear()
        self.transfers_tab.transfer_description_input.clear()
        self.transfers_tab.transfer_note_input.clear()
        self._refresh_transfer_form_choices()
        self.transfers_tab.transfer_from_account_combo.setCurrentIndex(0)
        self.transfers_tab.transfer_to_account_combo.setCurrentIndex(0)
        self.transfers_tab.transfer_save_button.setEnabled(True)
        self.transfers_tab.transfer_delete_button.setEnabled(False)

    def on_transfer_selection_changed(self) -> None:
        selections = self.transfers_tab.transfers_table.selectionModel().selectedRows()
        if not selections:
            self.editing_transfer_group_id = None
            self.editing_transfer_source_system = None
            return
        row = self.transfers_tab.transfers_model.row_dict(selections[0].row())
        if row is None:
            return
        self.editing_transfer_group_id = str(row.get("transfer_group_id") or "").strip() or None
        self.editing_transfer_source_system = str(row.get("source_system") or "").strip().lower()

        self.transfers_tab.transfer_date_input.setText(str(row.get("txn_date") or ""))
        self.transfers_tab.transfer_amount_input.setText(f"{int(row.get('amount_cents') or 0) / 100:.2f}")
        self.transfers_tab.transfer_description_input.setText(self._normalized_display_text(row.get("description")))
        self.transfers_tab.transfer_note_input.setText(self._normalized_display_text(row.get("note")))
        self._refresh_transfer_form_choices()
        self._combo_select_data(
            self.transfers_tab.transfer_from_account_combo,
            int(row.get("from_account_id") or 0) or None,
        )
        self._combo_select_data(
            self.transfers_tab.transfer_to_account_combo,
            int(row.get("to_account_id") or 0) or None,
        )

        is_manual = self.editing_transfer_source_system == "manual"
        for widget in (
            self.transfers_tab.transfer_date_input,
            self.transfers_tab.transfer_amount_input,
            self.transfers_tab.transfer_from_account_combo,
            self.transfers_tab.transfer_to_account_combo,
            self.transfers_tab.transfer_description_input,
            self.transfers_tab.transfer_note_input,
        ):
            widget.setEnabled(is_manual)
        self.transfers_tab.transfer_save_button.setEnabled(is_manual)
        self.transfers_tab.transfer_delete_button.setEnabled(is_manual)
        if not is_manual:
            self.statusBar().showMessage("Rule-based transfer selected (read-only).", 3000)

    def save_transfer_from_form(self) -> None:
        if len(self._transfer_account_choices()) < 2:
            QMessageBox.warning(
                self,
                "Transfer",
                "Define at least two active accounts in Settings before creating transfers.",
            )
            return
        try:
            transfer = self._build_transfer_input_from_form()
        except Exception as exc:  # noqa: BLE001
            QMessageBox.warning(self, "Transfer", str(exc))
            return

        editing_group_id = self.editing_transfer_group_id
        editing_source_system = str(self.editing_transfer_source_system or "").strip().lower()
        try:
            if editing_group_id and editing_source_system == "manual":
                updated = self.context.transactions_service.update_manual_transfer_group(editing_group_id, transfer)
                if not updated:
                    QMessageBox.warning(self, "Save Transfer", "Selected manual transfer no longer exists.")
                    self.refresh_transactions()
                    self.new_transfer_form()
                    return
                group_id = editing_group_id
                self.statusBar().showMessage("Transfer updated.", 3000)
                self.logger.info("Updated manual transfer group %s", group_id)
            elif editing_group_id and editing_source_system != "manual":
                QMessageBox.warning(self, "Save Transfer", "Rule-based transfers are read-only.")
                return
            else:
                group_id = self.context.transactions_service.add_transfer(transfer)
                self.statusBar().showMessage("Transfer added.", 3000)
                self.logger.info(
                    "Added transfer group %s (%s cents) from account %s to %s",
                    group_id,
                    transfer.amount_cents,
                    transfer.from_account_id,
                    transfer.to_account_id,
                )
        except Exception as exc:  # noqa: BLE001
            self.logger.error("Save transfer failed: %s", exc)
            QMessageBox.critical(self, "Save Transfer Failed", str(exc))
            return

        month_key = transfer.txn_date[:7]
        self._refresh_transactions_month_filter(preferred_month=month_key)
        self._refresh_transfers_month_filter(preferred_month=month_key)
        self._refresh_accounts_month_filter(preferred_month=month_key)
        self._suppress_selection_autoload = True
        try:
            self.refresh_transactions()
            self.refresh_transfers()
            self.refresh_accounts()
            self.refresh_dashboard()
            self._clear_transaction_form()
            self.new_transfer_form()
        finally:
            self._suppress_selection_autoload = False

    def delete_selected_transfer(self) -> None:
        transfer_group_id = str(self.editing_transfer_group_id or "").strip()
        if not transfer_group_id:
            QMessageBox.information(self, "Delete Transfer", "Select a transfer row to delete.")
            return
        source_system = str(self.editing_transfer_source_system or "").strip().lower()
        if source_system != "manual":
            QMessageBox.warning(self, "Delete Transfer", "Rule-based transfers cannot be deleted manually.")
            return
        answer = QMessageBox.question(
            self,
            "Delete Transfer",
            "Delete selected manual transfer?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if answer != QMessageBox.Yes:
            return
        deleted = self.context.transactions_service.delete_manual_transfer_group(transfer_group_id)
        if not deleted:
            QMessageBox.warning(self, "Delete Transfer", "Selected manual transfer no longer exists.")
            self.refresh_transactions()
            self.refresh_transfers()
            self.new_transfer_form()
            return
        self.refresh_transactions()
        self.refresh_transfers()
        self.refresh_accounts()
        self.refresh_dashboard()
        self.new_transfer_form()
        self.statusBar().showMessage("Transfer deleted.", 3000)
        self.logger.info("Deleted manual transfer group %s", transfer_group_id)

    def save_transaction(self, force_insert: bool = False) -> None:
        preferred_txn_month = (
            self.transactions_tab.month_filter.currentText().strip()
            or f"{self.transactions_view_year:04d}-{self.transactions_view_month:02d}"
        )
        preferred_accounts_month = (
            self.accounts_tab.month_filter.currentText().strip()
            or f"{self.accounts_view_year:04d}-{self.accounts_view_month:02d}"
        )
        editing_txn_id = None if force_insert else self.transactions_tab.editing_txn_id
        existing = None
        if editing_txn_id is not None:
            existing = self.context.transactions_service.get_transaction(editing_txn_id)
            if existing is None:
                QMessageBox.warning(self, "Save Transaction", "The selected transaction no longer exists.")
                self._clear_transaction_form()
                self.refresh_transactions()
                return
            if str(existing.get("txn_type") or "").strip().lower() == "transfer":
                QMessageBox.warning(
                    self,
                    "Save Transaction",
                    "Transfer actions are read-only on this tab. Use the Transfers tab.",
                )
                return

        try:
            txn = self._build_transaction_input(existing=existing)
            if editing_txn_id is None:
                txn_id = self.context.transactions_service.add_transaction(txn)
                self.logger.info("Added transaction %s", txn_id)
                self.statusBar().showMessage("Transaction added.", 3000)
            else:
                updated = self.context.transactions_service.update_transaction(editing_txn_id, txn)
                if not updated:
                    QMessageBox.warning(self, "Save Transaction", "No transaction was updated.")
                    return
                self.logger.info("Updated transaction %s", editing_txn_id)
                self.statusBar().showMessage("Transaction updated.", 3000)
        except ValueError as exc:
            QMessageBox.warning(self, "Save Transaction", str(exc))
            return
        except Exception as exc:
            self.logger.error("Save transaction failed: %s", exc)
            QMessageBox.critical(self, "Save Transaction Failed", str(exc))
            return

        self._refresh_transaction_form_choices()
        self._refresh_transactions_month_filter(preferred_month=preferred_txn_month)
        self._refresh_accounts_month_filter(preferred_month=preferred_accounts_month)
        self._suppress_selection_autoload = True
        try:
            self.refresh_transactions()
            self.refresh_accounts()
            self._clear_transaction_form()
        finally:
            self._suppress_selection_autoload = False

    def delete_selected_transaction(self) -> None:
        row = self._selected_transaction_row()
        if not row:
            QMessageBox.information(self, "Delete Transaction", "Select an expense or income row to delete.")
            return
        if str(row.get("txn_type") or "").strip().lower() == "transfer":
            QMessageBox.warning(
                self,
                "Delete Transaction",
                "Transfer actions cannot be deleted from Transactions. Use the Transfers tab.",
            )
            return

        txn_id = int(row["txn_id"])
        answer = QMessageBox.question(
            self,
            "Delete Transaction",
            f"Delete selected transaction #{txn_id}?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if answer != QMessageBox.Yes:
            return

        deleted = self.context.transactions_service.delete_transaction(txn_id)
        if not deleted:
            QMessageBox.warning(self, "Delete Transaction", "The selected transaction no longer exists.")
        else:
            self.logger.info("Deleted transaction %s", txn_id)
            self.statusBar().showMessage("Transaction deleted.", 3000)

        self._refresh_transactions_month_filter(
            preferred_month=f"{self.transactions_view_year}-{self.transactions_view_month:02d}"
        )
        self._refresh_accounts_month_filter(
            preferred_month=f"{self.transactions_view_year}-{self.transactions_view_month:02d}"
        )
        self.refresh_transactions()
        self.refresh_accounts()
        self._clear_transaction_form()

    @staticmethod
    def _build_bill_category_subtotals(rows: list[dict]) -> list[dict]:
        totals_by_category: dict[str, int] = {}
        for row in rows:
            category_name = str(row.get("category_name") or "Uncategorized").strip() or "Uncategorized"
            amount_cents = row.get("expected_amount_cents")
            cents = int(amount_cents) if amount_cents is not None else 0
            totals_by_category[category_name] = totals_by_category.get(category_name, 0) + cents

        subtotal_rows: list[dict] = []
        for category_name, cents in totals_by_category.items():
            subtotal_rows.append(
                {
                    "category_name": category_name,
                    "subtotal_display": f"${cents / 100:.2f}",
                }
            )
        subtotal_rows.sort(
            key=lambda row: (
                str(row.get("category_name") or "").strip().casefold(),
                str(row.get("category_name") or "").strip(),
            )
        )
        return subtotal_rows

    def refresh_bills(self) -> None:
        rows = self.context.bills_service.list_month_bills(
            sort_by=self.bills_sort_key,
            year=self.bills_view_year,
            month=self.bills_view_month,
        )
        self.bills_tab.model.replace_rows(rows)
        subtotal_rows = self._build_bill_category_subtotals(rows)
        self.bills_tab.category_totals_model.replace_rows(subtotal_rows)
        total_cents = sum(int(row.get("expected_amount_cents") or 0) for row in rows)
        self.bills_tab.total_subtotals_value_label.setText(f"${total_cents / 100:.2f}")
        subscriptions_rows = [
            row
            for row in rows
            if self._normalized_source_system(str(row.get("source_system") or "")) == "subtracker"
        ]
        subscriptions_count = len(subscriptions_rows)
        subscriptions_total_cents = sum(
            int(row.get("expected_amount_cents") or 0)
            for row in subscriptions_rows
        )
        self.bills_tab.subscriptions_totals_value_label.setText(
            f"{subscriptions_count} | ${subscriptions_total_cents / 100:.2f}"
        )
        self.logger.info(
            "Loaded %s bill occurrences for %s-%02d (sort=%s)",
            len(rows),
            self.bills_view_year,
            self.bills_view_month,
            self.bills_sort_key,
        )
        if self.budget_view_year == self.bills_view_year and self.budget_view_month == self.bills_view_month:
            self.refresh_budget_allocations()
        else:
            self.refresh_dashboard()

    def refresh_bills_for_selected_month(self) -> None:
        month_label = f"{self.bills_view_year:04d}-{self.bills_view_month:02d}"
        if self._is_bills_month_dirty(self.bills_view_year, self.bills_view_month):
            answer = QMessageBox.warning(
                self,
                "Refresh Bills",
                "Refreshing bills will replace all bill instances for "
                f"{month_label} with current global definitions.\n\n"
                "Any edits made to monthly bill instances for this month will be lost.\n\n"
                "Do you want to continue?",
                QMessageBox.Yes | QMessageBox.Cancel,
                QMessageBox.Cancel,
            )
            if answer != QMessageBox.Yes:
                self.statusBar().showMessage("Refresh Bills canceled.", 3000)
                return

        deleted, inserted = self.context.bills_service.regenerate_for_month(
            self.bills_view_year,
            self.bills_view_month,
        )
        self._clear_bills_month_dirty(self.bills_view_year, self.bills_view_month)
        self.refresh_bills()
        self.logger.info(
            "Refreshed bills for %s-%02d: replaced %s occurrences, generated %s from definitions.",
            self.bills_view_year,
            self.bills_view_month,
            deleted,
            inserted,
        )
        self.statusBar().showMessage(
            f"Refreshed bills for {month_label}: replaced {deleted}, generated {inserted}.",
            5000,
        )

    @staticmethod
    def _truncate_report_text(value: str, max_len: int) -> str:
        text = str(value or "").strip()
        if max_len <= 0:
            return ""
        if len(text) <= max_len:
            return text
        if max_len <= 3:
            return text[:max_len]
        return text[: max_len - 3] + "..."

    def _write_bills_report_pdf(self, save_path: Path, rows: list[dict], subtotal_rows: list[dict]) -> None:
        writer = QPdfWriter(str(save_path))
        writer.setResolution(120)
        page_layout = writer.pageLayout()
        page_layout.setPageSize(QPageSize(QPageSize.A4))
        page_layout.setOrientation(QPageLayout.Landscape)
        page_layout.setUnits(QPageLayout.Inch)
        page_layout.setMargins(QMarginsF(0.5, 0.5, 0.5, 0.5))
        writer.setPageLayout(page_layout)

        painter = QPainter()
        if not painter.begin(writer):
            raise OSError(f"Unable to create PDF at {save_path}")

        normal_font = QFont("Arial", 12)
        bold_font = QFont("Arial", 12)
        bold_font.setBold(True)
        painter.setFont(normal_font)

        metrics = painter.fontMetrics()
        line_height = metrics.height() + 6
        base_row_height = metrics.height() + 10

        # Compute the content bounds from the full page using explicit 0.5" margins
        # so all four sides are consistent across platforms/viewers.
        resolution = writer.resolution()
        margin_px = int(round(0.5 * resolution))
        full_rect = writer.pageLayout().fullRectPixels(resolution)
        content_rect = full_rect.adjusted(margin_px, margin_px, -margin_px, -margin_px)
        content_rect = content_rect.adjusted(1, 1, -1, -1)  # avoid edge clipping on borders
        left = content_rect.left()
        top = content_rect.top()
        bottom = content_rect.bottom()
        table_width = content_rect.width()

        y = top

        def ensure_space(required_height: int) -> None:
            nonlocal y
            if y + required_height <= bottom:
                return
            writer.newPage()
            painter.setFont(normal_font)
            y = top

        def draw_line(text: str, *, bold: bool = False) -> None:
            ensure_space(line_height)
            painter.setFont(bold_font if bold else normal_font)
            painter.drawText(
                QRect(left, y, table_width, line_height),
                int(Qt.AlignLeft | Qt.AlignVCenter),
                text,
            )
            painter.setFont(normal_font)
            nonlocal_y_increment()

        def nonlocal_y_increment(step: int = line_height) -> None:
            nonlocal y
            y += step

        def draw_table(
            *,
            title: str,
            columns: list[str],
            rows_data: list[list[str]],
            width_ratios: list[float],
            alignments: list[Qt.AlignmentFlag],
            wrap_columns: set[int] | None = None,
            bold_last_row: bool = False,
        ) -> None:
            nonlocal y
            if len(columns) != len(width_ratios) or len(columns) != len(alignments):
                raise ValueError("Columns, width ratios, and alignments must match.")
            wrap_columns = wrap_columns or set()

            widths = [int(table_width * ratio) for ratio in width_ratios]
            widths[-1] = table_width - sum(widths[:-1])

            def draw_header_row(section_title: str) -> None:
                nonlocal y
                draw_line(section_title, bold=True)
                ensure_space(base_row_height)
                x = left
                painter.setFont(bold_font)
                for idx, heading in enumerate(columns):
                    rect = QRect(x, y, widths[idx], base_row_height)
                    painter.fillRect(rect, QColor("#E5E7EB"))
                    painter.drawRect(rect)
                    painter.drawText(
                        rect.adjusted(6, 0, -6, 0),
                        int(Qt.AlignLeft | Qt.AlignVCenter),
                        heading,
                    )
                    x += widths[idx]
                painter.setFont(normal_font)
                y += base_row_height

            draw_header_row(title)
            for row_idx, row_cells in enumerate(rows_data):
                row_height = base_row_height
                for col_idx, cell in enumerate(row_cells):
                    if col_idx not in wrap_columns:
                        continue
                    probe_rect = QRect(0, 0, max(1, widths[col_idx] - 10), 10000)
                    wrapped_rect = painter.fontMetrics().boundingRect(
                        probe_rect,
                        int(Qt.TextWordWrap | Qt.AlignLeft | Qt.AlignTop),
                        str(cell or ""),
                    )
                    row_height = max(row_height, wrapped_rect.height() + 10)

                if y + row_height > bottom:
                    writer.newPage()
                    painter.setFont(normal_font)
                    y = top
                    draw_header_row(f"{title} (continued)")

                x = left
                is_last = row_idx == len(rows_data) - 1
                if bold_last_row and is_last:
                    painter.setFont(bold_font)
                for col_idx, cell in enumerate(row_cells):
                    rect = QRect(x, y, widths[col_idx], row_height)
                    painter.drawRect(rect)
                    text_rect = rect.adjusted(5, 4, -5, -4)
                    if col_idx in wrap_columns:
                        painter.drawText(
                            text_rect,
                            int(alignments[col_idx] | Qt.AlignTop | Qt.TextWordWrap),
                            str(cell or ""),
                        )
                    else:
                        cell_text = painter.fontMetrics().elidedText(
                            str(cell or ""),
                            Qt.ElideRight,
                            max(0, rect.width() - 10),
                        )
                        painter.drawText(
                            text_rect,
                            int(alignments[col_idx] | Qt.AlignVCenter),
                            cell_text,
                        )
                    x += widths[col_idx]
                painter.setFont(normal_font)
                y += row_height
            y += 6

        month_label = f"{self.bills_view_year:04d}-{self.bills_view_month:02d}"
        draw_line(f"BudgetPal Bills Report - {month_label}", bold=True)
        draw_line(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        y += 6

        # Section 1 first: Category Sub-Totals
        subtotal_table_rows: list[list[str]] = []
        total_cents = 0
        for row in subtotal_rows:
            category = str(row.get("category_name") or "").strip()
            subtotal_display = str(row.get("subtotal_display") or "$0.00").strip() or "$0.00"
            cents = int(
                (Decimal(subtotal_display.replace("$", "").replace(",", "")) * 100).quantize(
                    Decimal("1"),
                    rounding=ROUND_HALF_UP,
                )
            )
            total_cents += cents
            subtotal_table_rows.append([category, subtotal_display])
        subtotal_table_rows.append(["Total", f"${total_cents / 100:.2f}"])

        draw_table(
            title="Section 1: Category Sub-Totals",
            columns=["Category", "Subtotal"],
            rows_data=subtotal_table_rows,
            width_ratios=[0.72, 0.28],
            alignments=[Qt.AlignLeft, Qt.AlignRight],
            bold_last_row=True,
        )

        # Section 2 second: Bill details.
        bill_rows = [
            [
                str(row.get("category_name") or ""),
                str(row.get("name") or ""),
                str(row.get("payment_due") or ""),
                str(row.get("interval_display") or ""),
                str(row.get("amount_display") or ""),
                str(row.get("notes") or ""),
            ]
            for row in rows
        ]
        draw_table(
            title="Section 2: Bills and Subscriptions (Month Occurrences)",
            columns=["Category", "Name", "Payment Due", "Interval", "Amount", "Note"],
            rows_data=bill_rows,
            width_ratios=[0.14, 0.24, 0.13, 0.11, 0.10, 0.28],
            alignments=[Qt.AlignLeft, Qt.AlignLeft, Qt.AlignLeft, Qt.AlignLeft, Qt.AlignRight, Qt.AlignLeft],
            wrap_columns={5},
        )

        painter.end()

    @staticmethod
    def _rtf_escape(value: str) -> str:
        text = str(value or "")
        escaped: list[str] = []
        for ch in text:
            if ch == "\\":
                escaped.append(r"\\")
            elif ch == "{":
                escaped.append(r"\{")
            elif ch == "}":
                escaped.append(r"\}")
            elif ch == "\n":
                escaped.append(r"\line ")
            elif ord(ch) > 127:
                escaped.append(f"\\u{ord(ch)}?")
            else:
                escaped.append(ch)
        return "".join(escaped)

    @staticmethod
    def _rtf_cell_edges(total_width_twips: int, width_ratios: list[float]) -> list[int]:
        widths = [int(total_width_twips * ratio) for ratio in width_ratios]
        widths[-1] = total_width_twips - sum(widths[:-1])
        edges: list[int] = []
        running = 0
        for width in widths:
            running += max(1, int(width))
            edges.append(running)
        return edges

    def _build_rtf_table(
        self,
        *,
        title: str,
        columns: list[str],
        rows_data: list[list[str]],
        width_ratios: list[float],
        alignments: list[str],
        total_width_twips: int,
        bold_last_row: bool = False,
    ) -> list[str]:
        lines: list[str] = []
        lines.append(rf"\b {self._rtf_escape(title)}\b0\par")
        edges = self._rtf_cell_edges(total_width_twips, width_ratios)

        def row_rtf(cells: list[str], *, header: bool = False, bold_row: bool = False) -> str:
            parts: list[str] = [r"\trowd\trgaph108\trleft0"]
            for edge in edges:
                parts.append(rf"\cellx{edge}")
            for idx, raw in enumerate(cells):
                align = r"\ql"
                if idx < len(alignments) and alignments[idx] == "right":
                    align = r"\qr"
                content = self._rtf_escape(raw)
                prefix = align + r"\intbl "
                if header or bold_row:
                    prefix += r"\b "
                    suffix = r"\b0 \cell"
                else:
                    suffix = r"\cell"
                parts.append(f"{prefix}{content}{suffix}")
            parts.append(r"\row")
            return "".join(parts)

        lines.append(row_rtf(columns, header=True))
        for idx, row in enumerate(rows_data):
            is_last = idx == (len(rows_data) - 1)
            lines.append(row_rtf([str(cell or "") for cell in row], bold_row=bold_last_row and is_last))
        lines.append(r"\par")
        return lines

    def _write_bills_report_rtf(self, save_path: Path, rows: list[dict], subtotal_rows: list[dict]) -> None:
        # Letter landscape in twips: 11in x 8.5in -> 15840 x 12240
        # 0.5in margins all around -> 720 twips each side.
        paper_width = 15840
        paper_height = 12240
        margin = 720
        content_width = paper_width - (margin * 2)

        month_label = f"{self.bills_view_year:04d}-{self.bills_view_month:02d}"
        generated = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        subtotal_table_rows: list[list[str]] = []
        total_cents = 0
        for row in subtotal_rows:
            category = str(row.get("category_name") or "").strip()
            subtotal_display = str(row.get("subtotal_display") or "$0.00").strip() or "$0.00"
            cents = int(
                (Decimal(subtotal_display.replace("$", "").replace(",", "")) * 100).quantize(
                    Decimal("1"),
                    rounding=ROUND_HALF_UP,
                )
            )
            total_cents += cents
            subtotal_table_rows.append([category, subtotal_display])
        subtotal_table_rows.append(["Total", f"${total_cents / 100:.2f}"])

        bill_rows = [
            [
                str(row.get("category_name") or ""),
                str(row.get("name") or ""),
                str(row.get("payment_due") or ""),
                str(row.get("interval_display") or ""),
                str(row.get("amount_display") or ""),
                str(row.get("notes") or ""),
            ]
            for row in rows
        ]

        lines: list[str] = [
            r"{\rtf1\ansi\deff0",
            r"{\fonttbl{\f0 Arial;}}",
            rf"\paperw{paper_width}\paperh{paper_height}\landscape",
            rf"\margl{margin}\margr{margin}\margt{margin}\margb{margin}",
            r"\fs24",
            rf"\b {self._rtf_escape(f'BudgetPal Bills Report - {month_label}')}\b0\par",
            rf"Generated: {self._rtf_escape(generated)}\par\par",
        ]

        lines.extend(
            self._build_rtf_table(
                title="Section 1: Category Sub-Totals",
                columns=["Category", "Subtotal"],
                rows_data=subtotal_table_rows,
                width_ratios=[0.72, 0.28],
                alignments=["left", "right"],
                total_width_twips=content_width,
                bold_last_row=True,
            )
        )
        lines.extend(
            self._build_rtf_table(
                title="Section 2: Bills and Subscriptions (Month Occurrences)",
                columns=["Category", "Name", "Payment Due", "Interval", "Amount", "Note"],
                rows_data=bill_rows,
                width_ratios=[0.14, 0.24, 0.13, 0.11, 0.10, 0.28],
                alignments=["left", "left", "left", "left", "right", "left"],
                total_width_twips=content_width,
            )
        )
        lines.append("}")
        save_path.write_text("\n".join(lines), encoding="utf-8")

    def _write_bills_report_docx(self, save_path: Path, rows: list[dict], subtotal_rows: list[dict]) -> None:
        try:
            from docx import Document
            from docx.enum.section import WD_ORIENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
        except ImportError as exc:
            raise OSError(
                "python-docx is not installed. Install it with: pip install python-docx"
            ) from exc

        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        month_label = f"{self.bills_view_year:04d}-{self.bills_view_month:02d}"
        self._apply_docx_report_template(
            doc,
            report_name=f"Bills Report - {month_label}",
            generated_at=datetime.now(),
        )

        def set_cell_text(cell, text: str, *, bold: bool = False, align_right: bool = False) -> None:
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(text or ""))
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = bool(bold)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align_right else WD_ALIGN_PARAGRAPH.LEFT

        def add_table(title: str, headers: list[str], data_rows: list[list[str]], right_cols: set[int]) -> None:
            heading = doc.add_paragraph()
            heading_run = heading.add_run(title)
            heading_run.bold = True
            heading_run.font.name = "Arial"
            heading_run.font.size = Pt(12)

            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            for idx, header in enumerate(headers):
                set_cell_text(header_cells[idx], header, bold=True, align_right=False)

            for row in data_rows:
                cells = table.add_row().cells
                for idx, value in enumerate(row):
                    set_cell_text(
                        cells[idx],
                        str(value or ""),
                        align_right=(idx in right_cols),
                    )
            doc.add_paragraph("")

        subtotal_table_rows: list[list[str]] = []
        total_cents = 0
        for row in subtotal_rows:
            category = str(row.get("category_name") or "").strip()
            subtotal_display = str(row.get("subtotal_display") or "$0.00").strip() or "$0.00"
            cents = int(
                (Decimal(subtotal_display.replace("$", "").replace(",", "")) * 100).quantize(
                    Decimal("1"),
                    rounding=ROUND_HALF_UP,
                )
            )
            total_cents += cents
            subtotal_table_rows.append([category, subtotal_display])
        subtotal_table_rows.append(["Total", f"${total_cents / 100:.2f}"])

        add_table(
            "Section 1: Category Sub-Totals",
            ["Category", "Subtotal"],
            subtotal_table_rows,
            right_cols={1},
        )

        bill_rows = [
            [
                str(row.get("category_name") or ""),
                str(row.get("name") or ""),
                str(row.get("payment_due") or ""),
                str(row.get("interval_display") or ""),
                str(row.get("amount_display") or ""),
                str(row.get("notes") or ""),
            ]
            for row in rows
        ]
        add_table(
            "Section 2: Bills and Subscriptions (Month Occurrences)",
            ["Category", "Name", "Payment Due", "Interval", "Amount", "Note"],
            bill_rows,
            right_cols={4},
        )

        doc.save(str(save_path))

    def export_bills_report(self) -> None:
        start_dir = self._bills_report_dialog_start_dir()
        month_label = f"{self.bills_view_year:04d}-{self.bills_view_month:02d}"
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Bills Report (DOCX)",
            str(Path(start_dir) / f"budgetpal_bills_report_{month_label}.docx"),
            "Word Document (*.docx)",
        )
        if not file_path:
            return

        save_path = Path(file_path)
        if save_path.suffix.lower() != ".docx":
            save_path = save_path.with_suffix(".docx")
        if not save_path.stem.endswith(f"_{month_label}"):
            save_path = save_path.with_name(f"{save_path.stem}_{month_label}{save_path.suffix}")
        self._persist_last_bills_report_dir(save_path.parent)

        rows = self.context.bills_service.list_month_bills(
            sort_by=self.bills_sort_key,
            year=self.bills_view_year,
            month=self.bills_view_month,
        )
        subtotal_rows = self._build_bill_category_subtotals(rows)

        try:
            self._write_bills_report_docx(save_path, rows, subtotal_rows)
        except OSError as exc:
            self.logger.error("Bills report export failed: %s", exc)
            QMessageBox.critical(self, "Bills Report Export Failed", str(exc))
            return

        self.logger.info("Exported bills report to %s (%s rows, %s subtotals)", save_path, len(rows), len(subtotal_rows))
        self.statusBar().showMessage(f"Bills report saved: {save_path.name}", 5000)

    def _bills_report_dialog_start_dir(self) -> str:
        ui_settings = self.context.settings.get("ui", {})
        raw = str(ui_settings.get("last_bills_report_dir", "")).strip()
        if raw:
            candidate = Path(raw).expanduser()
            if candidate.is_file():
                candidate = candidate.parent
            if candidate.exists():
                return str(candidate)

        import_dir = str(ui_settings.get("last_import_dir", "")).strip()
        if import_dir:
            candidate = Path(import_dir).expanduser()
            if candidate.exists():
                return str(candidate)
        return str(Path.home())

    def _persist_last_bills_report_dir(self, directory: Path) -> None:
        try:
            resolved = directory.expanduser().resolve()
        except OSError:
            resolved = directory

        ui_settings = self.context.settings.setdefault("ui", {})
        new_value = str(resolved)
        if str(ui_settings.get("last_bills_report_dir", "")).strip() == new_value:
            return

        ui_settings["last_bills_report_dir"] = new_value
        try:
            get_settings_manager().save(self.context.settings)
        except OSError as exc:
            self.logger.error("Failed to save last bills report directory: %s", exc)

    def refresh_subscriptions(self) -> None:
        if not self.context.subscriptions_service:
            message = (
                "SubTracker DB path is not configured in budgetpal_config.json "
                "under subtracker.database_path"
            )
            self.logger.error(message)
            QMessageBox.warning(self, "SubTracker Not Configured", message)
            return

        month_label = f"{self.bills_view_year:04d}-{self.bills_view_month:02d}"
        if self._is_subscriptions_month_dirty(self.bills_view_year, self.bills_view_month):
            answer = QMessageBox.warning(
                self,
                "Refresh Subscriptions",
                "Refreshing subscriptions will replace subscription bill instances for "
                f"{month_label} with current SubTracker values.\n\n"
                "Any edits made to monthly subscription instances for this month will be lost.\n\n"
                "Do you want to continue?",
                QMessageBox.Yes | QMessageBox.Cancel,
                QMessageBox.Cancel,
            )
            if answer != QMessageBox.Yes:
                self.statusBar().showMessage("Refresh Subscriptions canceled.", 3000)
                return

        try:
            count = self.context.subscriptions_service.refresh_subtracker_bills(
                year=self.bills_view_year,
                month=self.bills_view_month,
            )
            mapping_errors = list(getattr(self.context.subscriptions_service, "last_mapping_errors", []))
            if mapping_errors:
                preview = "\n".join(mapping_errors[:5])
                if len(mapping_errors) > 5:
                    preview += f"\n...and {len(mapping_errors) - 5} more."
                self.logger.error(
                    "SubTracker category mapping fallback occurred for %s rows.",
                    len(mapping_errors),
                )
                QMessageBox.warning(
                    self,
                    "SubTracker Category Mapping Warning",
                    "Some SubTracker category IDs do not match BudgetPal categories.\n"
                    "Those subscriptions were assigned to Uncategorized.\n\n"
                    f"{preview}",
                )
            deleted, inserted = self.context.bills_service.regenerate_for_month(
                self.bills_view_year,
                self.bills_view_month,
                source_system="subtracker",
            )
            self._clear_bills_month_dirty(
                self.bills_view_year,
                self.bills_view_month,
                source_system="subtracker",
            )
            self.refresh_bills()
            self.logger.info("Refreshed %s active subscriptions from SubTracker", count)
            self.statusBar().showMessage(
                f"Refreshed subscriptions for {month_label}: definitions {count}, replaced {deleted}, generated {inserted}.",
                6000,
            )
        except SubTrackerIntegrationError as exc:
            self.logger.error("SubTracker integration failed: %s", exc)
            QMessageBox.critical(self, "SubTracker Integration Error", str(exc))

    def on_bill_definitions_changed(self) -> None:
        self.refresh_bills()
        self.statusBar().showMessage(
            "Bill definitions updated. Click Refresh Bills to regenerate the selected month.",
            6000,
        )

    def on_budget_definitions_changed(self) -> None:
        self.refresh_budget_allocations()
        self.statusBar().showMessage(
            "Budget category definitions updated. Click Refresh Category Budgets to regenerate the selected month.",
            6000,
        )

    def show_budget_definitions_dialog(self) -> None:
        if self.budget_definitions_dialog is None:
            self.budget_definitions_dialog = BudgetCategoryDefinitionsDialog(
                service=self.context.budget_allocations_service,
                categories_repo=self.context.categories_repo,
                logger=self.logger,
                on_changed=self.on_budget_definitions_changed,
                parent=self,
            )
            self.budget_definitions_dialog.setAttribute(Qt.WA_DeleteOnClose, False)

        self.budget_definitions_dialog.show()
        self.budget_definitions_dialog.raise_()
        self.budget_definitions_dialog.activateWindow()

    def show_bill_definitions_dialog(self) -> None:
        if self.bill_definitions_dialog is None:
            self.bill_definitions_dialog = RecurringDefinitionsDialog(
                bills_service=self.context.bills_service,
                categories_repo=self.context.categories_repo,
                logger=self.logger,
                on_changed=self.on_bill_definitions_changed,
                parent=self,
            )
            self.bill_definitions_dialog.setAttribute(Qt.WA_DeleteOnClose, False)

        self.bill_definitions_dialog.show()
        self.bill_definitions_dialog.raise_()
        self.bill_definitions_dialog.activateWindow()

    def on_income_definitions_changed(self) -> None:
        self.refresh_income()
        self.statusBar().showMessage(
            "Income definitions updated. Click Refresh Income to regenerate the selected month.",
            6000,
        )

    def show_income_definitions_dialog(self) -> None:
        if self.income_definitions_dialog is None:
            self.income_definitions_dialog = IncomeDefinitionsDialog(
                income_service=self.context.income_service,
                categories_repo=self.context.categories_repo,
                accounts_repo=self.context.accounts_repo,
                logger=self.logger,
                on_changed=self.on_income_definitions_changed,
                parent=self,
            )
            self.income_definitions_dialog.setAttribute(Qt.WA_DeleteOnClose, False)

        self.income_definitions_dialog.show()
        self.income_definitions_dialog.raise_()
        self.income_definitions_dialog.activateWindow()

    def show_sub_payments_dialog(self) -> None:
        if not self.context.subscription_payments_service:
            message = (
                "SubTracker DB path is not configured in budgetpal_config.json "
                "under subtracker.database_path"
            )
            QMessageBox.warning(self, "Sub Payments Unavailable", message)
            return

        if self.sub_payments_dialog is None:
            self.sub_payments_dialog = SubPaymentsDialog(
                service=self.context.subscription_payments_service,
                logger=self.logger,
                parent=self,
            )
            self.sub_payments_dialog.setAttribute(Qt.WA_DeleteOnClose, False)

        self.sub_payments_dialog.open_month(self.transactions_view_year, self.transactions_view_month)
        self.sub_payments_dialog.show()
        self.sub_payments_dialog.raise_()
        self.sub_payments_dialog.activateWindow()

    @staticmethod
    def _reports_period_label(year: int, month: int | None) -> str:
        if month is None:
            return f"{year:04d}"
        return f"{year:04d}-{month:02d}"

    @staticmethod
    def _slugify_report_name(name: str) -> str:
        slug = re.sub(r"[^a-zA-Z0-9]+", "_", str(name).strip().lower())
        return slug.strip("_") or "report"

    def _period_transactions(self, year: int) -> list[dict]:
        rows: list[dict] = []
        for m in range(1, 13):
            rows.extend(self.context.transactions_service.list_for_month(year=year, month=int(m), limit=50000))
        rows.sort(key=lambda row: (str(row.get("txn_date") or ""), int(row.get("txn_id") or 0)))
        return rows

    @staticmethod
    def _report_mode(report_key: str) -> str:
        row = report_type_lookup().get(str(report_key), {})
        mode = str(row.get("period_mode") or "").strip().lower()
        return mode if mode in {"year", "month"} else "month"

    @staticmethod
    def _clean_report_rows(rows: list[dict]) -> list[dict]:
        cleaned: list[dict] = []
        for row in rows:
            category_name = str(row.get("category_name") or "").strip()
            if not category_name:
                continue
            cleaned.append(dict(row))
        return cleaned

    def _account_month_metrics(self, account_id: int, year: int, month: int) -> dict[str, int]:
        beginning_balance = self.context.transactions_service.get_account_month_beginning_balance(
            year=year,
            month=month,
            account_id=account_id,
        )
        month_rows = self.context.transactions_service.list_account_ledger_for_month(
            year=year,
            month=month,
            account_id=account_id,
            include_prior_uncleared=False,
            limit=10000,
        )
        withdrawals = 0
        deposits = 0
        ending_balance = int(beginning_balance)
        for account_txn in month_rows:
            signed_amount = int(account_txn.get("amount_cents") or 0)
            if signed_amount > 0:
                deposits += signed_amount
            elif signed_amount < 0:
                withdrawals += abs(signed_amount)
            ending_balance += signed_amount
        return {
            "beginning": int(beginning_balance),
            "withdrawals": int(withdrawals),
            "deposits": int(deposits),
            "ending": int(ending_balance),
        }

    def _dashboard_account_status_rows(self, year: int, month: int) -> tuple[list[dict], int]:
        internal_account_rows = self.context.accounts_repo.list_active(include_external=False)
        account_status_rows: list[dict] = []
        total_internal_ending_balance_cents = 0
        for account_row in internal_account_rows:
            account_id = int(account_row["account_id"])
            account_label = str(account_row.get("name") or "").strip() or f"Account {account_id}"
            institution = str(account_row.get("institution_name") or "").strip()
            if institution:
                account_label = f"{institution} • {account_label}"

            metrics = self._account_month_metrics(account_id=account_id, year=year, month=month)
            total_internal_ending_balance_cents += int(metrics["ending"])
            account_status_rows.append(
                {
                    "account_name": account_label,
                    "beginning_display": self._format_currency_balance(int(metrics["beginning"])),
                    "activity_display": (
                        f"W: {self._format_currency(int(metrics['withdrawals']))} / "
                        f"D: {self._format_currency(int(metrics['deposits']))}"
                    ),
                    "ending_display": self._format_currency_balance(int(metrics["ending"])),
                }
            )
        account_status_rows.sort(key=lambda item: str(item.get("account_name") or "").casefold())
        return account_status_rows, int(total_internal_ending_balance_cents)

    def _build_tax_preparation_report(self, year: int) -> tuple[str, str]:
        period_label = self._reports_period_label(year, None)
        title = f"Tax Preparation Report - {period_label}"
        income_transactions: list[dict] = []
        expense_transactions: list[dict] = []
        income_by_category: dict[str, int] = {}
        expense_by_category: dict[str, int] = {}

        for row in self._period_transactions(year):
            if not bool(row.get("tax_deductible")):
                continue
            txn_type = str(row.get("txn_type") or "").strip().lower()
            category = self._normalized_display_text(row.get("category_name")) or "Uncategorized"
            amount_cents = abs(int(row.get("amount_cents") or 0))
            if txn_type == "income":
                income_transactions.append(dict(row))
                income_by_category[category] = income_by_category.get(category, 0) + amount_cents
            elif txn_type == "expense":
                expense_transactions.append(dict(row))
                expense_by_category[category] = expense_by_category.get(category, 0) + amount_cents

        income_transactions.sort(
            key=lambda row: (
                self._normalized_display_text(row.get("category_name")) or "Uncategorized",
                str(row.get("txn_date") or ""),
                int(row.get("txn_id") or 0),
            )
        )
        expense_transactions.sort(
            key=lambda row: (
                self._normalized_display_text(row.get("category_name")) or "Uncategorized",
                str(row.get("txn_date") or ""),
                int(row.get("txn_id") or 0),
            )
        )

        lines = [title, ""]
        lines.append("Section 1A: Taxable Income Transactions")
        lines.append("Category | Date | Description | Account | Amount")
        if not income_transactions:
            lines.append("None |  |  |  | $0.00")
        else:
            for row in income_transactions:
                category = self._normalized_display_text(row.get("category_name")) or "Uncategorized"
                lines.append(
                    f"{category} | "
                    f"{row.get('txn_date') or ''} | "
                    f"{self._normalized_display_text(row.get('description'))} | "
                    f"{self._normalized_display_text(row.get('account_name'))} | "
                    f"{self._format_currency(abs(int(row.get('amount_cents') or 0)))}"
                )
        lines.append("")

        lines.append("Section 1B: Taxable Income by Category")
        lines.append("Category | Subtotal")
        income_total = 0
        if not income_by_category:
            lines.append("None | $0.00")
        else:
            for category in sorted(income_by_category, key=str.casefold):
                subtotal = int(income_by_category[category])
                income_total += subtotal
                lines.append(f"{category} | {self._format_currency(subtotal)}")
        lines.append(f"Total | {self._format_currency(income_total)}")
        lines.append("")

        lines.append("Section 2A: Tax-Deductible Expense Transactions")
        lines.append("Category | Date | Description | Account | Amount")
        if not expense_transactions:
            lines.append("None |  |  |  | $0.00")
        else:
            for row in expense_transactions:
                category = self._normalized_display_text(row.get("category_name")) or "Uncategorized"
                lines.append(
                    f"{category} | "
                    f"{row.get('txn_date') or ''} | "
                    f"{self._normalized_display_text(row.get('description'))} | "
                    f"{self._normalized_display_text(row.get('account_name'))} | "
                    f"{self._format_currency(abs(int(row.get('amount_cents') or 0)))}"
                )
        lines.append("")

        lines.append("Section 2B: Tax-Deductible Expenses by Category")
        lines.append("Category | Subtotal")
        expense_total = 0
        if not expense_by_category:
            lines.append("None | $0.00")
        else:
            for category in sorted(expense_by_category, key=str.casefold):
                subtotal = int(expense_by_category[category])
                expense_total += subtotal
                lines.append(f"{category} | {self._format_currency(subtotal)}")
        lines.append(f"Total | {self._format_currency(expense_total)}")
        return title, "\n".join(lines)

    def _build_dashboard_report(self, year: int, month: int) -> tuple[str, str]:
        period_label = self._reports_period_label(year, month)
        title = f"Dashboard Report - {period_label}"
        snapshot = self._compute_dashboard_snapshot_for_month(year, month)
        account_rows, total_internal_ending = self._dashboard_account_status_rows(year, month)

        lines = [title, ""]
        lines.append("Section 1: Budget Allocation Details - Expenses")
        lines.append("Category | Planned | Actual | Diff")
        for row in self._clean_report_rows(list(snapshot["expense_rows"])):
            lines.append(
                f"{row.get('category_name') or ''} | "
                f"{row.get('planned_display') or ''} | "
                f"{row.get('actual_display') or ''} | "
                f"{row.get('diff_display') or ''}"
            )
        lines.append("")

        lines.append("Section 2: Budget Allocation Details - Income")
        lines.append("Category | Planned | Actual | Diff")
        for row in self._clean_report_rows(list(snapshot["income_rows"])):
            lines.append(
                f"{row.get('category_name') or ''} | "
                f"{row.get('planned_display') or ''} | "
                f"{row.get('actual_display') or ''} | "
                f"{row.get('diff_display') or ''}"
            )
        lines.append("")

        lines.append("Section 3: Account Status")
        lines.append("Account | Beginning | Activity (W/D) | Ending")
        for row in account_rows:
            lines.append(
                f"{row.get('account_name') or ''} | "
                f"{row.get('beginning_display') or ''} | "
                f"{row.get('activity_display') or ''} | "
                f"{row.get('ending_display') or ''}"
            )
        lines.append(
            f"Total Internal Ending Balance |  |  | {self._format_currency_balance(total_internal_ending)}"
        )
        return title, "\n".join(lines)

    def _build_annual_account_status_report(self, year: int) -> tuple[str, str]:
        title = f"Annual Account Status Report - {self._reports_period_label(year, None)}"
        accounts = self.context.accounts_repo.list_active(include_external=False)

        lines = [title, "", "Section 1: Annual Account Status", "Account | Beginning | Activity (W/D) | Ending"]
        total_beginning = 0
        total_withdrawals = 0
        total_deposits = 0
        total_ending = 0
        if not accounts:
            lines.append("No internal accounts defined. | $0.00 | W: $0.00 / D: $0.00 | $0.00")
            return title, "\n".join(lines)

        for account in sorted(
            accounts,
            key=lambda row: (
                str(row.get("institution_name") or "").casefold(),
                str(row.get("name") or "").casefold(),
                int(row.get("account_id") or 0),
            ),
        ):
            beginning_sum = 0
            ending_sum = 0
            withdrawals_sum = 0
            deposits_sum = 0
            account_id = int(account.get("account_id") or 0)
            for month in range(1, 13):
                metrics = self._account_month_metrics(account_id=account_id, year=year, month=month)
                beginning_sum += int(metrics["beginning"])
                ending_sum += int(metrics["ending"])
                withdrawals_sum += int(metrics["withdrawals"])
                deposits_sum += int(metrics["deposits"])
            total_beginning += beginning_sum
            total_ending += ending_sum
            total_withdrawals += withdrawals_sum
            total_deposits += deposits_sum
            account_name = str(account.get("name") or "").strip() or f"Account {account_id}"
            institution = str(account.get("institution_name") or "").strip()
            if institution:
                account_name = f"{institution} • {account_name}"
            lines.append(
                f"{account_name} | "
                f"{self._format_currency(beginning_sum)} | "
                f"W: {self._format_currency(withdrawals_sum)} / D: {self._format_currency(deposits_sum)} | "
                f"{self._format_currency(ending_sum)}"
            )

        lines.append(
            f"Total | {self._format_currency(total_beginning)} | "
            f"W: {self._format_currency(total_withdrawals)} / D: {self._format_currency(total_deposits)} | "
            f"{self._format_currency(total_ending)}"
        )
        return title, "\n".join(lines)

    def _build_report_content(self, report_key: str, year: int, month: int | None) -> tuple[str, str]:
        key = str(report_key or "").strip()
        if key == "tax_preparation":
            return self._build_tax_preparation_report(year)
        if key == "dashboard_monthly":
            if month is None:
                raise ValueError("Dashboard Report requires a Year and Month selection.")
            return self._build_dashboard_report(year, int(month))
        if key == "annual_account_status":
            return self._build_annual_account_status_report(year)
        raise ValueError(f"Unsupported report key: {key}")

    def _reports_export_start_dir(self) -> str:
        ui_settings = self.context.settings.get("ui", {})
        raw = str(ui_settings.get("last_reports_export_dir", "")).strip()
        if raw:
            candidate = Path(raw).expanduser()
            if candidate.exists() and candidate.is_dir():
                return str(candidate)
        import_dir = str(ui_settings.get("last_import_dir", "")).strip()
        if import_dir:
            candidate = Path(import_dir).expanduser()
            if candidate.exists() and candidate.is_dir():
                return str(candidate)
        return str(Path.home())

    def _persist_last_reports_export_dir(self, directory: Path) -> None:
        try:
            resolved = directory.expanduser().resolve()
        except OSError:
            resolved = directory
        ui_settings = self.context.settings.setdefault("ui", {})
        new_value = str(resolved)
        if str(ui_settings.get("last_reports_export_dir", "")).strip() == new_value:
            return
        ui_settings["last_reports_export_dir"] = new_value
        try:
            get_settings_manager().save(self.context.settings)
        except OSError as exc:
            self.logger.error("Failed to persist reports export directory: %s", exc)

    def _restore_reports_table_column_widths(self) -> None:
        def _as_int(value, fallback: int) -> int:
            try:
                return int(value)
            except (TypeError, ValueError):
                return int(fallback)

        ui_settings = self.context.settings.get("ui", {})
        reports_ui = ui_settings.get("reports", {})
        columns = reports_ui.get("columns", {})
        report_width = _as_int(
            columns.get("report", self.reports_tab.DEFAULT_REPORT_COLUMN_WIDTH),
            self.reports_tab.DEFAULT_REPORT_COLUMN_WIDTH,
        )
        description_width = _as_int(
            columns.get("description", self.reports_tab.DEFAULT_DESCRIPTION_COLUMN_WIDTH),
            self.reports_tab.DEFAULT_DESCRIPTION_COLUMN_WIDTH,
        )
        self.reports_tab.set_column_widths(report_width, description_width)

    def _persist_reports_table_column_widths(self) -> None:
        report_width, description_width = self.reports_tab.column_widths()
        ui_settings = self.context.settings.setdefault("ui", {})
        reports_ui = ui_settings.setdefault("reports", {})
        columns = reports_ui.setdefault("columns", {})
        new_columns = {
            "report": int(report_width),
            "description": int(description_width),
        }
        try:
            existing_report = int(
                columns.get("report", self.reports_tab.DEFAULT_REPORT_COLUMN_WIDTH)
            )
        except (TypeError, ValueError):
            existing_report = self.reports_tab.DEFAULT_REPORT_COLUMN_WIDTH
        try:
            existing_description = int(
                columns.get("description", self.reports_tab.DEFAULT_DESCRIPTION_COLUMN_WIDTH)
            )
        except (TypeError, ValueError):
            existing_description = self.reports_tab.DEFAULT_DESCRIPTION_COLUMN_WIDTH
        if (
            existing_report == new_columns["report"]
            and existing_description == new_columns["description"]
        ):
            return
        reports_ui["columns"] = new_columns
        try:
            get_settings_manager().save(self.context.settings)
        except OSError as exc:
            self.logger.error("Failed to persist reports table column widths: %s", exc)
            return
        self.logger.info(
            "Saved Reports table column widths: report=%s, description=%s",
            new_columns["report"],
            new_columns["description"],
        )

    @staticmethod
    def _add_docx_page_number(paragraph) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        paragraph.alignment = 2  # right
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    @staticmethod
    def _add_docx_horizontal_divider(paragraph) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        if not paragraph.text:
            paragraph.add_run(" ")
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    @staticmethod
    def _set_docx_cell_bottom_border(cell, *, size: int = 10, color: str = "000000") -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        bottom = tc_borders.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            tc_borders.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(int(size)))
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), color)

    @classmethod
    def _apply_docx_report_template(cls, doc, report_name: str, generated_at: datetime) -> None:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt

        section = doc.sections[0]
        margin = Inches(0.5)
        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = margin
        section.bottom_margin = margin

        header_paragraph = (
            section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        )
        header_paragraph.text = ""
        cls._add_docx_page_number(header_paragraph)

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(12)

        brand_table = doc.add_table(rows=1, cols=3)
        brand_table.autofit = False
        brand_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        usable_width = section.page_width - section.left_margin - section.right_margin
        left_width = int(usable_width * 0.18)
        right_width = int(usable_width * 0.18)
        center_width = int(usable_width - left_width - right_width)
        brand_table.columns[0].width = left_width
        brand_table.columns[1].width = center_width
        brand_table.columns[2].width = right_width
        row = brand_table.rows[0]
        row.cells[0].width = left_width
        row.cells[1].width = center_width
        row.cells[2].width = right_width
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        # Force table width to span from margin to margin.
        tbl_pr = brand_table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(int(usable_width / 635)))  # EMU -> twips
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:type"), "dxa")
        tbl_ind.set(qn("w:w"), "0")

        logo_cell = row.cells[0]
        logo_cell.text = ""
        logo_p = logo_cell.paragraphs[0]
        logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_p.paragraph_format.space_before = Pt(0)
        logo_p.paragraph_format.space_after = Pt(0)
        logo_path = BudgetPalPathRegistry.logo_image_file()
        if logo_path and logo_path.exists():
            logo_p.add_run().add_picture(str(logo_path), width=Inches(0.95))
        else:
            logo_p.add_run("BudgetPal")

        title_cell = row.cells[1]
        title_cell.text = ""
        brand_p = title_cell.paragraphs[0]
        brand_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        brand_p.paragraph_format.space_before = Pt(0)
        brand_p.paragraph_format.space_after = Pt(0)
        brand_run = brand_p.add_run("BudgetPal")
        brand_run.bold = True
        brand_run.font.name = "Arial"
        brand_run.font.size = Pt(16)
        report_p = title_cell.add_paragraph()
        report_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_p.paragraph_format.space_before = Pt(0)
        report_p.paragraph_format.space_after = Pt(0)
        report_run = report_p.add_run(report_name)
        report_run.bold = True
        report_run.font.name = "Arial"
        report_run.font.size = Pt(14)

        date_cell = row.cells[2]
        date_cell.text = ""
        date_p = date_cell.paragraphs[0]
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_p.paragraph_format.space_before = Pt(0)
        date_p.paragraph_format.space_after = Pt(0)
        date_run = date_p.add_run(generated_at.strftime("%Y-%m-%d"))
        date_run.font.name = "Arial"
        date_run.font.size = Pt(12)

        cls._set_docx_cell_bottom_border(logo_cell, size=12, color="000000")
        cls._set_docx_cell_bottom_border(title_cell, size=12, color="000000")
        cls._set_docx_cell_bottom_border(date_cell, size=12, color="000000")
        spacer = doc.add_paragraph("")
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(8)

    @classmethod
    def _write_report_docx(cls, save_path: Path, title: str, body_text: str) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        doc = Document()
        cls._apply_docx_report_template(
            doc,
            report_name=title,
            generated_at=datetime.now(),
        )

        lines = body_text.splitlines()
        if lines and lines[0].strip() == title.strip():
            lines = lines[1:]
            if lines and not lines[0].strip():
                lines = lines[1:]

        def _is_divider_line(line: str) -> bool:
            stripped = str(line or "").strip()
            return bool(stripped) and set(stripped) <= {"-", "|", " "}

        def _looks_currency_or_numeric(value: str) -> bool:
            text = str(value or "").strip().replace(",", "")
            if not text:
                return False
            if text.endswith("%"):
                text = text[:-1]
            if text.startswith(("+$", "-$", "$")):
                text = text.replace("$", "")
            elif text.startswith(("+", "-")):
                text = text[1:]
            try:
                float(text)
                return True
            except ValueError:
                return False

        idx = 0
        while idx < len(lines):
            line = str(lines[idx] or "")
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                idx += 1
                continue

            if "|" in line and not _is_divider_line(line):
                raw_rows: list[list[str]] = []
                while idx < len(lines):
                    current = str(lines[idx] or "")
                    if not current.strip():
                        break
                    if "|" not in current:
                        break
                    if _is_divider_line(current):
                        idx += 1
                        continue
                    raw_rows.append([part.strip() for part in current.split("|")])
                    idx += 1

                if raw_rows:
                    column_count = max(len(row) for row in raw_rows)
                    rows = [row + [""] * (column_count - len(row)) for row in raw_rows]
                    table = doc.add_table(rows=len(rows), cols=column_count)
                    table.style = "Table Grid"
                    for r_idx, row_values in enumerate(rows):
                        for c_idx, cell_text in enumerate(row_values):
                            cell = table.cell(r_idx, c_idx)
                            cell.text = cell_text
                            paragraph = cell.paragraphs[0]
                            align_right = c_idx > 0 and _looks_currency_or_numeric(cell_text)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align_right else WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.name = "Arial"
                                run.font.size = Pt(12)
                                if r_idx == 0:
                                    run.bold = True
                    doc.add_paragraph("")
                continue

            para = doc.add_paragraph(line)
            if stripped.lower().startswith("section "):
                for run in para.runs:
                    run.bold = True
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(12)
            idx += 1
        doc.save(str(save_path))

    def run_selected_report(self) -> None:
        selected_reports = self.reports_tab.selected_reports()
        if not selected_reports:
            QMessageBox.information(self, "Reports", "Select one or more reports to run.")
            return

        try:
            year = int(self.reports_tab.year_picker.currentText())
        except ValueError:
            QMessageBox.warning(self, "Reports", "Select a valid year before running reports.")
            return
        month_text = self.reports_tab.month_picker.currentText().strip()
        month_value = int(month_text) if month_text else None

        missing_month_reports: list[str] = []
        outputs: list[tuple[str, str, str, str]] = []
        for report_key, report_label in selected_reports:
            mode = self._report_mode(report_key)
            if mode == "month":
                if month_value is None:
                    missing_month_reports.append(report_label)
                    continue
                effective_month = int(month_value)
            else:
                effective_month = None
            try:
                title, body = self._build_report_content(report_key, year, effective_month)
            except ValueError as exc:
                QMessageBox.warning(self, "Reports", f"{report_label}: {exc}")
                continue
            period_label = self._reports_period_label(year, effective_month)
            outputs.append((report_label, title, body, period_label))

        if missing_month_reports:
            QMessageBox.warning(
                self,
                "Reports",
                "The following report(s) require Year/Month selection:\n\n"
                + "\n".join(missing_month_reports),
            )
        if not outputs:
            return

        export_dir_raw = QFileDialog.getExistingDirectory(
            self,
            "Select Reports Export Directory",
            self._reports_export_start_dir(),
        )
        if not export_dir_raw:
            return
        export_dir = Path(export_dir_raw)
        self._persist_last_reports_export_dir(export_dir)

        exported_files: list[Path] = []
        for report_label, title, body, period_label in outputs:
            slug = self._slugify_report_name(report_label)
            output_file = export_dir / f"budgetpal_{slug}_{period_label}.docx"
            self._write_report_docx(output_file, title, body)
            exported_files.append(output_file)

        if self.reports_tab.preview_after_export_checkbox.isChecked():
            open_failures: list[str] = []
            for output_file in exported_files:
                if not QDesktopServices.openUrl(QUrl.fromLocalFile(str(output_file))):
                    open_failures.append(output_file.name)
            if open_failures:
                QMessageBox.warning(
                    self,
                    "Preview Failed",
                    "Some exported reports could not be opened automatically:\n\n"
                    + "\n".join(open_failures),
                )

        self.logger.info(
            "Exported %s report(s) to %s",
            len(exported_files),
            export_dir,
        )
        self.statusBar().showMessage(
            f"Exported {len(exported_files)} report(s) to {export_dir}",
            6000,
        )

    @staticmethod
    def _extract_log_level(message: str) -> str:
        parts = [part.strip() for part in message.split("|", 3)]
        if len(parts) >= 2 and parts[1]:
            return parts[1].upper()
        return "INFO"

    def append_log_message(self, message: str) -> None:
        level = self._extract_log_level(message)
        color_hex = self.LOG_LEVEL_COLORS.get(level, self.LOG_LEVEL_COLORS["DEFAULT"])
        self.log_area.setTextColor(QColor(color_hex))
        self.log_area.append(message)
        self.log_area.setTextColor(QColor(self.LOG_LEVEL_COLORS["DEFAULT"]))

    def _clear_log(self) -> None:
        self.log_area.clear()
        self.logger.info("Activity log cleared by user")

    def show_settings_dialog(self) -> None:
        dialog = SettingsDialog(
            settings=self.context.settings,
            categories_repo=self.context.categories_repo,
            accounts_repo=self.context.accounts_repo,
            backup_now_callback=self.backup_database_now,
            clear_monthly_data_callback=self.clear_monthly_transactions_and_instances_now,
            validate_subtracker_categories_callback=self.validate_subtracker_categories_now,
            export_definitions_callback=self.export_global_definitions_now,
            import_definitions_callback=self.import_global_definitions_now,
            generate_transactions_template_callback=self.generate_transactions_template_now,
            logger=self.logger,
            parent=self,
        )
        accepted = bool(dialog.exec())

        if dialog.categories_dirty:
            self._refresh_transaction_form_choices()
            self._refresh_budget_form_choices()
            self._refresh_bill_form_choices()
            self.refresh_transactions()
            self.refresh_budget_allocations()
            self.refresh_bills()
            self.logger.info("Categories updated from Settings dialog")

        if dialog.accounts_dirty:
            self._refresh_transaction_form_choices()
            self._refresh_income_form_choices()
            self._refresh_accounts_month_filter(
                preferred_month=f"{self.accounts_view_year}-{self.accounts_view_month:02d}"
            )
            self.refresh_transactions()
            self.refresh_income()
            self.refresh_accounts()
            self.logger.info("Accounts updated from Settings dialog")

        if not accepted:
            return

        new_settings = dialog.settings_value()
        old_db_path = str(self.context.settings.get("database", {}).get("path", "")).strip()
        old_subtracker_db_path = str(
            self.context.settings.get("subtracker", {}).get("database_path", "")
        ).strip()
        old_backup_cfg = dict(self.context.settings.get("backup", {}))
        old_transfer_rules = list(self.context.settings.get("transfers", {}).get("rules", []))
        old_categories_export_dir = str(
            self.context.settings.get("ui", {}).get("last_categories_export_dir", "")
        ).strip()
        old_definitions_export_dir = str(
            self.context.settings.get("ui", {}).get("last_definitions_export_dir", "")
        ).strip()
        old_definitions_import_dir = str(
            self.context.settings.get("ui", {}).get("last_definitions_import_dir", "")
        ).strip()
        old_definitions_import_file = str(
            self.context.settings.get("ui", {}).get("last_definitions_import_file", "")
        ).strip()
        old_window_cfg = self.context.settings.get("ui", {}).get("window", {})
        old_window_width = int(old_window_cfg.get("width", 1240))
        old_window_height = int(old_window_cfg.get("height", 820))
        old_logging_cfg = dict(self.context.settings.get("logging", {}))

        settings_mgr = get_settings_manager()
        try:
            settings_mgr.save(new_settings)
        except OSError as exc:
            self.logger.error("Failed to save settings: %s", exc)
            QMessageBox.critical(
                self,
                "Settings Save Error",
                "BudgetPal could not write the settings file.\n\n"
                f"{exc}\n\n"
                "Check file permissions and try again.",
            )
            return
        self.context.refresh_settings(new_settings)

        window_cfg = new_settings.get("ui", {}).get("window", {})
        configured_width = int(window_cfg.get("width", 1240))
        configured_height = int(window_cfg.get("height", 820))
        if configured_width != old_window_width or configured_height != old_window_height:
            self.resize(max(800, int(round(configured_width * 1.25))), configured_height)

        level_name = str(new_settings.get("logging", {}).get("level", "INFO")).upper()
        self.logger.setLevel(getattr(logging, level_name, logging.INFO))

        new_db_path = str(new_settings.get("database", {}).get("path", "")).strip()
        new_subtracker_db_path = str(new_settings.get("subtracker", {}).get("database_path", "")).strip()
        new_backup_cfg = dict(new_settings.get("backup", {}))
        new_transfer_rules = list(new_settings.get("transfers", {}).get("rules", []))
        restart_needed = old_db_path != new_db_path
        subtracker_changed = old_subtracker_db_path != new_subtracker_db_path

        self.logger.info("Settings saved to config file")
        if old_db_path != new_db_path:
            self.logger.info("Setting changed: database.path -> %s", new_db_path)
        if old_subtracker_db_path != new_subtracker_db_path:
            self.logger.info("Setting changed: subtracker.database_path -> %s", new_subtracker_db_path)
        if old_backup_cfg != new_backup_cfg:
            self.logger.info(
                "Setting changed: backup -> directory=%s, base_name=%s",
                str(new_backup_cfg.get("directory", "")),
                str(new_backup_cfg.get("base_name", "")),
            )
        if old_transfer_rules != new_transfer_rules:
            self.logger.info(
                "Setting changed: transfers.rules -> %s rule(s)",
                len(new_transfer_rules),
            )
        new_categories_export_dir = str(
            new_settings.get("ui", {}).get("last_categories_export_dir", "")
        ).strip()
        if old_categories_export_dir != new_categories_export_dir:
            self.logger.info(
                "Setting changed: ui.last_categories_export_dir -> %s",
                new_categories_export_dir,
            )
        new_definitions_export_dir = str(
            new_settings.get("ui", {}).get("last_definitions_export_dir", "")
        ).strip()
        if old_definitions_export_dir != new_definitions_export_dir:
            self.logger.info(
                "Setting changed: ui.last_definitions_export_dir -> %s",
                new_definitions_export_dir,
            )
        new_definitions_import_dir = str(
            new_settings.get("ui", {}).get("last_definitions_import_dir", "")
        ).strip()
        if old_definitions_import_dir != new_definitions_import_dir:
            self.logger.info(
                "Setting changed: ui.last_definitions_import_dir -> %s",
                new_definitions_import_dir,
            )
        new_definitions_import_file = str(
            new_settings.get("ui", {}).get("last_definitions_import_file", "")
        ).strip()
        if old_definitions_import_file != new_definitions_import_file:
            self.logger.info(
                "Setting changed: ui.last_definitions_import_file -> %s",
                new_definitions_import_file,
            )
        if configured_width != old_window_width or configured_height != old_window_height:
            self.logger.info(
                "Setting changed: ui.window -> %sx%s",
                configured_width,
                configured_height,
            )
        if str(old_logging_cfg.get("level", "INFO")).upper() != level_name:
            self.logger.info("Setting changed: logging.level -> %s", level_name)
        new_max_bytes = int(new_settings.get("logging", {}).get("max_bytes", 1_000_000))
        new_backup_count = int(new_settings.get("logging", {}).get("backup_count", 5))
        if int(old_logging_cfg.get("max_bytes", 1_000_000)) != new_max_bytes:
            self.logger.info("Setting changed: logging.max_bytes -> %s", new_max_bytes)
        if int(old_logging_cfg.get("backup_count", 5)) != new_backup_count:
            self.logger.info("Setting changed: logging.backup_count -> %s", new_backup_count)
        if restart_needed:
            self.statusBar().showMessage(
                "Settings saved. Restart required for BudgetPal DB path changes.",
                7000,
            )
            QMessageBox.information(
                self,
                "Settings Saved",
                "Settings saved to config.\n\n"
                "Restart BudgetPal to apply the new BudgetPal database path.",
            )
        else:
            self.statusBar().showMessage("Settings saved.", 4000)

        if subtracker_changed:
            if self.sub_payments_dialog is not None:
                self.sub_payments_dialog.close()
                self.sub_payments_dialog = None
            self.logger.info("SubTracker path updated; integration binding refreshed")

    @staticmethod
    def _sanitize_backup_base_name(base_name: str) -> str:
        raw = str(base_name or "").strip()
        if not raw:
            return "budgetpal_backup"
        cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", raw).strip("._-")
        return cleaned or "budgetpal_backup"

    def clear_monthly_transactions_and_instances_now(
        self,
        scope: str,
        year: int | None,
        month: int | None,
    ) -> dict[str, int]:
        # Keep global definitions intact; only wipe month-instance and transaction data.
        normalized_scope = str(scope or "year_month").strip().lower()
        if normalized_scope not in {"year_month", "year", "all"}:
            raise ValueError("Invalid clear scope. Expected year_month, year, or all.")
        if normalized_scope in {"year_month", "year"} and year is None:
            raise ValueError("Year is required for selected clear scope.")
        if normalized_scope == "year_month" and month is None:
            raise ValueError("Month is required for Year/Month clear scope.")
        if month is not None and (month < 1 or month > 12):
            raise ValueError("Month must be in the range 1-12.")

        year_str = f"{int(year):04d}" if year is not None else None
        month_key = f"{int(year):04d}-{int(month):02d}" if year is not None and month is not None else None
        txn_where = "1=1"
        txn_params: list[str] = []
        ym_where = "1=1"
        ym_params: list[str] = []

        if normalized_scope == "year_month":
            txn_where = "(import_period_key = ? OR substr(txn_date, 1, 7) = ?)"
            txn_params = [str(month_key), str(month_key)]
            ym_where = "year = ? AND month = ?"
            ym_params = [str(year), str(month)]
        elif normalized_scope == "year":
            txn_where = "(substr(coalesce(import_period_key, ''), 1, 4) = ? OR substr(txn_date, 1, 4) = ?)"
            txn_params = [str(year_str), str(year_str)]
            ym_where = "year = ?"
            ym_params = [str(year)]

        deleted_counts: dict[str, int] = {
            "transactions": 0,
            "bill_occurrences": 0,
            "income_occurrences": 0,
            "budget_lines": 0,
            "budget_months": 0,
            "account_month_settings": 0,
            "sub_payment_mappings": 0,
            "bills_month_settings": 0,
        }

        with self.context.db.connection() as conn:
            table_names = {
                str(row["name"])
                for row in conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()
            }

            if "bill_occurrences" in table_names:
                cur = conn.execute(
                    f"DELETE FROM bill_occurrences WHERE {ym_where}",
                    ym_params,
                )
                deleted_counts["bill_occurrences"] = int(cur.rowcount or 0)

            if "income_occurrences" in table_names:
                cur = conn.execute(
                    f"DELETE FROM income_occurrences WHERE {ym_where}",
                    ym_params,
                )
                deleted_counts["income_occurrences"] = int(cur.rowcount or 0)

            if "budget_months" in table_names and "budget_lines" in table_names:
                cur = conn.execute(
                    f"""
                    DELETE FROM budget_lines
                    WHERE budget_month_id IN (
                        SELECT budget_month_id
                        FROM budget_months
                        WHERE {ym_where}
                    )
                    """,
                    ym_params,
                )
                deleted_counts["budget_lines"] = int(cur.rowcount or 0)

                cur = conn.execute(
                    f"DELETE FROM budget_months WHERE {ym_where}",
                    ym_params,
                )
                deleted_counts["budget_months"] = int(cur.rowcount or 0)

            if "account_month_settings" in table_names:
                cur = conn.execute(
                    f"DELETE FROM account_month_settings WHERE {ym_where}",
                    ym_params,
                )
                deleted_counts["account_month_settings"] = int(cur.rowcount or 0)

            if "bills_month_settings" in table_names:
                cur = conn.execute(
                    f"DELETE FROM bills_month_settings WHERE {ym_where}",
                    ym_params,
                )
                deleted_counts["bills_month_settings"] = int(cur.rowcount or 0)

            if "sub_payment_mappings" in table_names and "transactions" in table_names:
                cur = conn.execute(
                    f"""
                    DELETE FROM sub_payment_mappings
                    WHERE txn_id IN (
                        SELECT txn_id
                        FROM transactions
                        WHERE {txn_where}
                    )
                    """,
                    txn_params,
                )
                deleted_counts["sub_payment_mappings"] = int(cur.rowcount or 0)

            if "transactions" in table_names:
                cur = conn.execute(
                    f"DELETE FROM transactions WHERE {txn_where}",
                    txn_params,
                )
                deleted_counts["transactions"] = int(cur.rowcount or 0)

        self._bills_dirty_by_month.clear()
        self._subscriptions_dirty_by_month.clear()
        self._income_dirty_by_month.clear()
        self._budget_dirty_by_month.clear()

        preferred_month = date.today().strftime("%Y-%m")
        self._refresh_dashboard_month_filter(preferred_month=preferred_month)
        self.on_dashboard_month_changed(preferred_month)

        total_deleted = int(sum(deleted_counts.values()))
        deleted_counts["total_deleted"] = total_deleted
        scope_label = "all"
        if normalized_scope == "year":
            scope_label = str(year_str)
        elif normalized_scope == "year_month":
            scope_label = str(month_key)
        self.logger.warning(
            "Cleared monthly transactions and instances from Settings (scope=%s total_deleted=%s).",
            scope_label,
            total_deleted,
        )
        self.statusBar().showMessage(
            f"Monthly data cleared ({total_deleted} rows deleted).",
            6000,
        )
        return deleted_counts

    def validate_subtracker_categories_now(self) -> dict[str, object]:
        service = self.context.subscriptions_service
        if service is None:
            raise ValueError(
                "SubTracker DB path is not configured in budgetpal_config.json "
                "under subtracker.database_path"
            )

        result = service.validate_category_mapping()
        total = int(result.get("total_subscriptions", 0))
        issue_count = int(result.get("issue_count", 0))
        if issue_count == 0:
            self.logger.info(
                "SubTracker category validation passed (total subscriptions=%s).",
                total,
            )
        else:
            self.logger.warning(
                "SubTracker category validation found %s issue(s) across %s subscriptions.",
                issue_count,
                total,
            )
        return result

    def backup_database_now(self, directory: Path, base_name: str) -> Path:
        target_dir = Path(directory).expanduser()
        if not target_dir.exists() or not target_dir.is_dir():
            raise OSError("Backup location is not reachable.")
        safe_base = self._sanitize_backup_base_name(base_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = target_dir / f"{safe_base}_{timestamp}.sqlite"

        source = sqlite3.connect(str(self.context.db.db_path))
        dest = sqlite3.connect(str(output_path))
        try:
            source.backup(dest)
        finally:
            source.close()
            dest.close()

        pruned = self._prune_old_backups(target_dir, safe_base, keep_count=self.BACKUP_KEEP_COUNT)

        self.logger.info("Database backup created: %s", output_path)
        if pruned > 0:
            self.logger.info(
                "Pruned %s old backup file(s); keeping latest %s.",
                pruned,
                self.BACKUP_KEEP_COUNT,
            )
        self.statusBar().showMessage(f"Backup complete: {output_path.name}", 4000)
        return output_path

    def _prune_old_backups(self, directory: Path, safe_base: str, keep_count: int) -> int:
        if keep_count < 1:
            keep_count = 1
        pattern = f"{safe_base}_*.sqlite"
        backups = [path for path in directory.glob(pattern) if path.is_file()]
        if len(backups) <= keep_count:
            return 0

        backups.sort(
            key=lambda path: path.stat().st_mtime,
            reverse=True,
        )
        removed = 0
        for old_path in backups[keep_count:]:
            try:
                old_path.unlink()
            except OSError as exc:
                self.logger.warning("Could not prune old backup %s: %s", old_path, exc)
                continue
            removed += 1
        return removed

    def export_global_definitions_now(self, directory: Path) -> list[Path]:
        target_dir = Path(directory).expanduser()
        if not target_dir.exists() or not target_dir.is_dir():
            raise OSError("Definitions export location is not reachable.")
        outputs = self.reporting_service.export_global_definitions(target_dir)
        self.logger.info(
            "Global definitions exported to %s (%s files).",
            target_dir,
            len(outputs),
        )
        self.statusBar().showMessage(
            f"Exported global definitions to {target_dir}",
            5000,
        )
        return outputs

    @staticmethod
    def _normalize_template_header_text(value: object) -> str:
        if value is None:
            return ""
        return str(value).strip().lower()

    @classmethod
    def _infer_transactions_template_section(
        cls,
        ws,
        header_row: int,
        header_col: int,
        *,
        max_scan_cols: int,
    ) -> str:
        best_label = "unknown"
        best_score: float | None = None
        for row in range(max(1, header_row - 12), header_row):
            for col in range(1, max_scan_cols + 1):
                token = cls._normalize_template_header_text(ws.cell(row=row, column=col).value)
                if not token:
                    continue
                label = None
                if "expense" in token:
                    label = "expense"
                elif "income" in token:
                    label = "income"
                if label is None:
                    continue
                vertical_distance = header_row - row
                horizontal_distance = abs(header_col - col)
                # Heavier weight on vertical distance so nearby section headers win.
                score = float(vertical_distance * 10 + horizontal_distance)
                if best_score is None or score < best_score:
                    best_score = score
                    best_label = label
        return best_label

    @classmethod
    def _find_transactions_template_column_ranges(
        cls,
        ws,
        target_header: str,
        section: str | None = None,
    ) -> list[str]:
        try:
            from openpyxl.utils import get_column_letter
        except ModuleNotFoundError as exc:
            raise RuntimeError(
                "openpyxl is required for transaction template generation. "
                "Install it with: pip install openpyxl"
            ) from exc

        target_token = cls._normalize_template_header_text(target_header)
        section_token = cls._normalize_template_header_text(section) if section else ""
        max_scan_rows = min(int(ws.max_row), 120)
        max_scan_cols = max(12, int(ws.max_column))
        header_pattern = ["date", "amount", "description", "category"]
        optional_headers = {
            "account",
            "subscription",
            "tax",
            "type",
            "payment type",
            "payment_type",
            "note",
            "notes",
        }
        ranges: list[str] = []

        for row in range(1, max_scan_rows + 1):
            for col in range(1, max_scan_cols - 2):
                headers = [
                    cls._normalize_template_header_text(ws.cell(row=row, column=col + idx).value)
                    for idx in range(4)
                ]
                if headers != header_pattern:
                    continue

                section_name = cls._infer_transactions_template_section(
                    ws,
                    row,
                    col,
                    max_scan_cols=max_scan_cols,
                )
                header_cols: dict[str, int] = {"category": col + 3}
                next_col = col + 4
                while next_col <= max_scan_cols:
                    header_name = cls._normalize_template_header_text(
                        ws.cell(row=row, column=next_col).value
                    )
                    if not header_name:
                        break
                    if header_name in optional_headers:
                        if header_name not in header_cols:
                            header_cols[header_name] = next_col
                        next_col += 1
                        continue
                    break

                target_col = header_cols.get(target_token)
                if target_col is not None and (not section_token or section_name == section_token):
                    start_row = row + 1
                    end_row = max(200, start_row)
                    col_letter = get_column_letter(target_col)
                    ranges.append(f"{col_letter}{start_row}:{col_letter}{end_row}")

        seen: set[str] = set()
        deduped: list[str] = []
        for entry in ranges:
            if entry in seen:
                continue
            seen.add(entry)
            deduped.append(entry)
        return deduped

    @classmethod
    def _find_transactions_template_account_ranges(cls, ws) -> list[str]:
        return cls._find_transactions_template_column_ranges(ws, "account")

    def generate_transactions_template_now(self, output_path: Path) -> Path:
        try:
            from openpyxl import load_workbook
            from openpyxl.worksheet.datavalidation import DataValidation
        except ModuleNotFoundError as exc:
            raise RuntimeError(
                "openpyxl is required for transaction template generation. "
                "Install it with: pip install openpyxl"
            ) from exc

        source_template = BudgetPalPathRegistry.transactions_template_file()
        if source_template is None:
            raise FileNotFoundError(
                "Could not locate the bundled transaction template file."
            )

        account_rows = self.context.accounts_repo.list_active()
        alias_rows: list[tuple[str, str, int]] = []
        alias_seen: set[str] = set()
        duplicates: list[str] = []
        for row in account_rows:
            alias = str(row.get("name") or "").strip()
            if not alias:
                continue
            alias_key = alias.casefold()
            if alias_key in alias_seen:
                duplicates.append(alias)
                continue
            alias_seen.add(alias_key)
            institution = str(row.get("institution_name") or "").strip()
            account_id = int(row.get("account_id") or 0)
            alias_rows.append((alias, institution, account_id))

        if duplicates:
            duplicate_text = ", ".join(sorted(set(duplicates), key=str.casefold))
            raise ValueError(
                "Cannot generate template because account aliases are not unique.\n\n"
                f"Duplicate aliases: {duplicate_text}\n\n"
                "Fix account aliases in Settings > Accounts first."
            )
        if not alias_rows:
            raise ValueError(
                "No active account aliases are available. "
                "Add at least one account in Settings > Accounts first."
            )

        alias_rows.sort(key=lambda item: (item[1].casefold(), item[0].casefold(), item[2]))
        aliases = [item[0] for item in alias_rows]
        def _normalized_category_names(rows: list[dict]) -> list[str]:
            deduped: dict[str, str] = {}
            for row in rows:
                name = str(row.get("name") or "").strip()
                if not name:
                    continue
                key = name.casefold()
                if key not in deduped:
                    deduped[key] = name
            return [deduped[key] for key in sorted(deduped.keys())]

        expense_categories = _normalized_category_names(
            self.context.categories_repo.list_active(category_type="expense")
        )
        income_categories = _normalized_category_names(
            self.context.categories_repo.list_active(category_type="income")
        )
        if not expense_categories:
            raise ValueError(
                "No active expense categories are available. "
                "Define categories in Settings > Definitions first."
            )
        if not income_categories:
            raise ValueError(
                "No active income categories are available. "
                "Define categories in Settings > Definitions first."
            )

        workbook = load_workbook(source_template)
        if "Transactions" not in workbook.sheetnames:
            raise ValueError("Template workbook is missing worksheet 'Transactions'.")
        ws = workbook["Transactions"]
        account_ranges = self._find_transactions_template_account_ranges(ws)
        expense_category_ranges = self._find_transactions_template_column_ranges(
            ws,
            "category",
            section="expense",
        )
        income_category_ranges = self._find_transactions_template_column_ranges(
            ws,
            "category",
            section="income",
        )
        if not account_ranges:
            raise ValueError(
                "Template workbook is missing Account columns in Transactions sections."
            )
        if not expense_category_ranges:
            raise ValueError(
                "Template workbook is missing Expense Category columns in Transactions sections."
            )
        if not income_category_ranges:
            raise ValueError(
                "Template workbook is missing Income Category columns in Transactions sections."
            )

        # Compatibility-first approach:
        # - Numbers and Google Sheets commonly preserve Excel list validations better
        #   when they are inline literal lists.
        # - If aliases are too many/long for inline validation, fall back to a helper range.
        helper_name = "__bp_template_lists"
        helper_ws = workbook[helper_name] if helper_name in workbook.sheetnames else workbook.create_sheet(
            title=helper_name
        )
        helper_ws.delete_cols(1, helper_ws.max_column)
        helper_ws.delete_rows(1, helper_ws.max_row)
        helper_ws.sheet_state = "visible"

        def _build_list_formula(values: list[str], helper_col_index: int) -> tuple[str, str]:
            csv_values = ",".join(values)
            can_use_inline = ("," not in "".join(values)) and (len(csv_values) <= 240)
            if can_use_inline:
                return f'"{csv_values}"', "inline"
            for index, value in enumerate(values, start=1):
                helper_ws.cell(row=index, column=helper_col_index, value=value)
            helper_col_letter = chr(ord("A") + helper_col_index - 1)
            return (
                f"='{helper_name}'!${helper_col_letter}$1:${helper_col_letter}${len(values)}",
                "range",
            )

        accounts_formula, accounts_validation_mode = _build_list_formula(aliases, 1)
        expense_category_formula, expense_validation_mode = _build_list_formula(expense_categories, 2)
        income_category_formula, income_validation_mode = _build_list_formula(income_categories, 3)

        existing_validations = list(ws.data_validations.dataValidation)

        def _upsert_list_validation(
            ranges: list[str],
            formula: str,
            error_title: str,
            error_text: str,
        ) -> None:
            normalized_ranges = [entry.replace(" ", "") for entry in ranges]
            for validation_range in normalized_ranges:
                updated_existing = False
                for dv in existing_validations:
                    if str(dv.type or "").lower() != "list":
                        continue
                    sqref_tokens = [token.replace(" ", "") for token in str(dv.sqref).split()]
                    if validation_range not in sqref_tokens:
                        continue
                    dv.formula1 = formula
                    dv.allow_blank = True
                    dv.showDropDown = False
                    dv.errorTitle = error_title
                    dv.error = error_text
                    updated_existing = True
                if updated_existing:
                    continue
                dv = DataValidation(type="list", formula1=formula, allow_blank=True)
                dv.showDropDown = False
                dv.errorTitle = error_title
                dv.error = error_text
                dv.add(validation_range)
                ws.add_data_validation(dv)

        _upsert_list_validation(
            account_ranges,
            accounts_formula,
            "Invalid Account Alias",
            "Select an account alias from the list.",
        )
        _upsert_list_validation(
            expense_category_ranges,
            expense_category_formula,
            "Invalid Expense Category",
            "Select an expense category from the list.",
        )
        _upsert_list_validation(
            income_category_ranges,
            income_category_formula,
            "Invalid Income Category",
            "Select an income category from the list.",
        )

        target = Path(output_path).expanduser()
        target.parent.mkdir(parents=True, exist_ok=True)
        workbook.save(target)
        self.logger.info(
            "Generated transactions template at %s using %s account alias(es) (%s), "
            "%s expense categories (%s), and %s income categories (%s).",
            target,
            len(aliases),
            accounts_validation_mode,
            len(expense_categories),
            expense_validation_mode,
            len(income_categories),
            income_validation_mode,
        )
        self.statusBar().showMessage(
            f"Generated transactions template: {target.name}",
            5000,
        )
        return target.resolve()

    def import_global_definitions_now(
        self,
        definition_type: str,
        csv_path: Path,
    ) -> dict[str, int | str]:
        source_path = Path(csv_path).expanduser()
        if not source_path.exists() or not source_path.is_file():
            raise OSError("Definitions import file is not reachable.")

        result = self.reporting_service.import_global_definitions(definition_type, source_path)

        self._refresh_transaction_form_choices()
        self._refresh_transfer_form_choices()
        self._refresh_bill_form_choices()
        self._refresh_income_form_choices()
        self._refresh_budget_form_choices()
        self._refresh_accounts_month_filter(
            preferred_month=f"{self.accounts_view_year}-{self.accounts_view_month:02d}"
        )
        self.refresh_transactions()
        self.refresh_transfers()
        self.refresh_accounts()
        self.refresh_bills()
        self.refresh_income()
        self.refresh_budget_allocations()
        self.refresh_dashboard()
        self.logger.info(
            "Imported global definitions from %s (type=%s inserted=%s updated=%s skipped_blank=%s)",
            source_path,
            result.get("definition_type", definition_type),
            result.get("inserted", 0),
            result.get("updated", 0),
            result.get("skipped_blank", 0),
        )
        self.statusBar().showMessage(
            f"Imported {result.get('definition_type', definition_type)} definitions from {source_path.name}",
            5000,
        )
        return result

    def _backup_database_on_exit(self) -> None:
        backup_cfg = self.context.settings.get("backup", {})
        directory_raw = str(backup_cfg.get("directory", "")).strip()
        if not directory_raw:
            self.logger.info("Exit backup skipped: backup directory not configured.")
            return
        base_name = str(backup_cfg.get("base_name", "budgetpal_backup")).strip() or "budgetpal_backup"
        try:
            self.backup_database_now(Path(directory_raw), base_name)
            self.logger.info("Exit backup completed successfully.")
        except Exception as exc:  # noqa: BLE001
            self.logger.error("Exit backup failed: %s", exc)

    def closeEvent(self, event) -> None:  # noqa: N802
        self._persist_reports_table_column_widths()
        self._backup_database_on_exit()
        super().closeEvent(event)

    def show_about_dialog(self) -> None:
        info = load_build_info()
        QMessageBox.about(
            self,
            "About BudgetPal",
            "BudgetPal\n"
            "Spend well!.\n\n"
            f"Version: {info.version}\n"
            f"Commit: {info.commit}\n"
            f"Built (UTC): {info.built_at_utc}",
        )

    def show_help(self) -> None:
        try:
            opened = self.help_service.open_main_help()
        except FileNotFoundError as exc:
            QMessageBox.warning(
                self,
                "Help Not Found",
                str(exc),
            )
            return
        except Exception as exc:  # noqa: BLE001
            QMessageBox.warning(
                self,
                "Help Error",
                f"Could not open help: {exc}",
            )
            return

        if not opened:
            try:
                path = self.help_service.get_topic_path("index")
                QMessageBox.information(
                    self,
                    "Help",
                    "Could not open the browser automatically.\n"
                    f"Open this file manually:\n{path}",
                )
            except Exception:  # noqa: BLE001
                pass
